#!/usr/bin/env python3
"""
Generator profesionálního Word dokumentu s 18 otázkami pro Karla Šmída
k forwarduje projektantům (SMASH, TeAnau, TUSPO) k vyřešení.

Source: inputs/meta/vyjasneni_queue.json
Output: outputs/Otazky_pro_Karla_a_projektanty_2026-05-18.docx

Audience: Karel Šmíd (zhotovitel — chytrý řemeslník, ne projektant ani rozpočtář).
Styl: čistá čeština, žádné anglické technické termíny, slovní hodnocení místo
'severity', vysvětlení kontextu pro každou otázku.
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Cm, Pt, RGBColor

PROJ = Path(__file__).resolve().parent.parent
QUEUE_JSON = PROJ / "inputs" / "meta" / "vyjasneni_queue.json"
OUT = PROJ / "outputs" / "Otazky_pro_Karla_a_projektanty_2026-05-18.docx"

# ───────────────────────────────────────────────────────────────────────────
# Otázky — friendly Czech rewrite pro Karel'a (per #id z queue)

OTAZKY = {
    1: {
        "title": "Vybrat rozsah rozpočtu — agregovaný / položkový / hybrid",
        "komu": "Investor Mgr. Jindřich Volný",
        "komu_email": "JindraVolny@seznam.cz",
        "vaznost": "URGENTNÍ",
        "o_co_jde": (
            "Pan Volný v emailu zmínil, že rozpočtář by mu mohl udělat „nějaký agregovaný“ "
            "rozpočet ze stavebního povolení. Před tím, než dokončíme cenovou nabídku, "
            "potřebujeme jednoznačné rozhodnutí, kterou variantu rozsahu si vybírá:\n"
            "  • Varianta A — agregovaný (souhrny po kapitolách, cca 30 řádků), nejlevnější\n"
            "  • Varianta B — položkový (každá položka samostatně, cca 187 řádků), nejdražší\n"
            "  • Varianta C — hybrid (HSV položkově, PSV a TZB agregovaně, kompromis)"
        ),
        "co_je_treba": (
            "Karle, prosím přepošli toto Volnému s žádostí o vybrání jedné varianty. "
            "Bez potvrzení nemůžeme dokončit cenovou nabídku."
        ),
        "co_mame": (
            "Momentálně rozpočet máme připravený ve všech 3 variantách (Excel s listy "
            "Var_A_Agregovany, Var_C_Hybrid, Var_B_Polozkovy). Po výběru jen vyplníme "
            "jednotkové ceny pro vybranou variantu."
        ),
        "vliv": "Blokuje finalizaci cenové nabídky. Bez výběru nelze pokračovat.",
    },
    2: {
        "title": "Parkovací stání u Dvořákovy — kam patří?",
        "komu": "Investor Volný + Architekt SMASH (Ing. arch. Smolka, ČKA 05394)",
        "komu_email": "JindraVolny@seznam.cz; SMASH architekti",
        "vaznost": "DŮLEŽITÉ",
        "o_co_jde": (
            "Souhrnná technická zpráva (kap. m.5) uvádí samostatně povolovanou akci "
            "„parkovacích stání a vjezdu na pozemek ze severní strany z ulice Dvořákova“. "
            "Rozhodnutí o vjezdu vydáno 17. 9. 2025, spis 0003300/25/Pne. "
            "Ale není jasné, zda samotná parkovací stání jsou součástí stavební akce "
            "260217 (sklad + parking), nebo zda jsou třetím samostatným stavebním objektem."
        ),
        "co_je_treba": (
            "Karle, prosím ověř u Bc. Jiřího Šmída nebo přímo u SMASHE, jak jsou parkovací "
            "stání vedena. Pokud jsou součástí 260217, zůstávají v jednom rozpočtu se skladem. "
            "Pokud jsou samostatný objekt, budou v rozpočtu odděleny."
        ),
        "co_mame": (
            "Momentálně použito: parkovací stání = součást 260217 (sklad+parking) "
            "podle architekonické technické zprávy D.1.1.00, kapitola 1.3 Provozní řešení."
        ),
        "vliv": (
            "Pouze členění rozpočtu do objektů. Celková cena se nemění, ale prezentace "
            "investorovi může vypadat jinak."
        ),
    },
    3: {
        "title": "Chybí výkresy ve výrobním stupni (DPS)",
        "komu": "Architekt SMASH (Ing. arch. Marek Smolka)",
        "komu_email": "SMASH architekti s.r.o., ČKA 05394",
        "vaznost": "DŮLEŽITÉ",
        "o_co_jde": (
            "Architekt zpracoval pouze dokumentaci pro stavební povolení (DSP). "
            "Dokumentace pro provádění stavby (DPS) nebyla zpracována. To znamená, že "
            "nemáme k dispozici tabulku skladeb podlah a stěn s konkrétními F-kódy, "
            "výpisy oken, dveří a klempířiny, ani detaily kotvení a lemování. "
            "Naštěstí v DXF výkresech jsme našli embedded tabulku místností s přesnými "
            "plochami pro 25 místností a 7 typů oken s přesnými rozměry — to nám "
            "dovolilo udělat rozpočet podstatně přesnější, než je u DSP obvyklé."
        ),
        "co_je_treba": (
            "Pokud chcete přesný položkový rozpočet (přesnost ±5 %), je třeba objednat "
            "u SMASHE doplnění DPS. Orientační cena: 150–200 tis. Kč.\n"
            "Pokud DPS nedoplní, rozpočet zůstává v rozsahu DSP s přesností ±10–15 % "
            "a s jasným označením, kde jsou hodnoty převzaty z výkresů a kde "
            "doplněny obvyklými hodnotami podle ČSN."
        ),
        "co_mame": (
            "Z DXF se podařilo vytáhnout: 25 místností s přesnými plochami (celkem "
            "217,5 m² vs technická zpráva uvádí 219,3 m² — přesnost 0,82 %), 16 oken "
            "v 7 typech (přesné rozměry každého typu z CAD bloků), vnější obvod domu "
            "38,70 m, 12 skladeb vrstev z TZ (podlahy, stropy, stěny, krov, fasáda)."
        ),
        "vliv": (
            "Pokud zůstane DSP, doporučujeme v cenové nabídce Volnému uvést rezervu "
            "10–15 % pro variabilitu finálních výměr u PSV (povrchové úpravy, výplně "
            "otvorů, podlahy)."
        ),
    },
    4: {
        "title": "Skladby podlah — přesné vrstvy a výrobce  [RESOLVED 2026-05-18]",
        "komu": "Architekt SMASH",
        "komu_email": "SMASH architekti s.r.o.",
        "vaznost": "INFORMATIVNÍ",
        "status": "RESOLVED",
        "status_note": (
            "VYŘEŠENO 2026-05-18 (Path C Tier 2): 13 S-codes decodováno z TZ řez A-A legendy + "
            "cross-referenced s DXF S-code markers per element location. Sheet 8 Var_E_Skladby_Vrstev "
            "obsahuje plné composition vrstev pro všech 13 S-codes (S01–S12b). Výrobce systému zůstává "
            "otázkou pro investora (PSV Knauf / Rigips / Fermacell — rozhodnutí o dodavateli)."
        ),
        "o_co_jde": (
            "Technická zpráva uvádí materiály a tloušťky vrstev pro 3 typy skladeb "
            "podlah (na terénu, suchá nad trámovým stropem, mokrá nad ocelobetonem). "
            "Ale konkrétní výrobce systému (Knauf / Rigips / Fermacell pro suchou "
            "skladbu) ani přesné F-kódy jednotlivých vrstev nejsou specifikovány."
        ),
        "co_je_treba": (
            "Pokud má SMASH zpracovanou tabulku skladeb (typicky list s kódem F-01, "
            "F-02, atd. pro každou skladbu), prosíme o její dodání. "
            "Pokud ne, doporučujeme stavbu konzultovat s konkrétním dodavatelem "
            "systému (Knauf nebo Fermacell pro suchou skladbu) před zahájením."
        ),
        "co_mame": (
            "Momentálně použito (s explicit označením „převzato z technické zprávy“):\n"
            "  • Podlaha 1.NP terén: ŽB deska + hydroizolace + EPS 150 tl. 120 mm "
            "(λ=0,035) + potěr s kari + nivelace + nášlap (vinyl/dlažba)\n"
            "  • Suchá podlaha: prkenný záklop + zásyp liapor/keramzit + kročejová "
            "minerální deska + Fermacell 25 mm + nášlap\n"
            "  • Mokrá podlaha 3.NP: ocelobeton + kročejová EPS + potěr s kari + nášlap"
        ),
        "vliv": (
            "Materiálové ceny mezi výrobci se mohou lišit o ±20 %. Pro přesnost "
            "rozpočtu by bylo užitečné mít rozhodnutí o konkrétním systému."
        ),
    },
    5: {
        "title": "Výpisy oken, dveří, klempířiny, zámečnictví, truhlářství  [PARTIALLY RESOLVED 2026-05-18]",
        "komu": "Architekt SMASH",
        "komu_email": "SMASH architekti s.r.o.",
        "vaznost": "STŘEDNĚ DŮLEŽITÉ",
        "status": "PARTIALLY_RESOLVED",
        "status_note": (
            "ČÁSTEČNĚ VYŘEŠENO 2026-05-18 (Path C Tier 4 — INSERT block extrakce z DXF): "
            "Z DXF se podařilo vytáhnout 7 unique window types s exact bbox rozměry — to nahrazuje "
            "typický výpis oken pro DSP. Klempířina 173,8 m total z DXF MA_klempíř + SM__klempířina "
            "layers, per-typ split (atika / okap / svod / parapet) je v Sheet 8 Var_E + Var_B detail. "
            "ZBÝVÁ: TZ silent o materiálu / fabrice klempířiny — dotaz na výrobce zůstává. "
            "Dveře vnitřní stále odhad ±20 % (ARS uvádí 'DTD laminované' bez ks/typů)."
        ),
        "o_co_jde": (
            "Technická zpráva uvádí obecné popisy: „plastová okna trojsklem Uw=0,85, "
            "barva dle investora; venkovní žaluzie v kastlíku s purenitovou izolací; "
            "vnitřní dveře DTD laminované; truhlářské stupně schodiště; ocelové "
            "zábradlí svařované z jeklů“. Ale konkrétní výpisy s počty kusů a "
            "přesnými rozměry nejsou. Naštěstí z DXF jsme získali 7 typů oken s "
            "přesnými rozměry (např. okno 2.NP 1160×1480 mm × 4 kusy)."
        ),
        "co_je_treba": (
            "Pokud má SMASH vypracované výpisy (typicky tabulky pro každý profesní "
            "typ — okna, dveře, klempířina, zámeč. prvky, truhlářské prvky), prosíme "
            "o jejich dodání. To nám umožní upřesnit počty a rozměry."
        ),
        "co_mame": (
            "Okna: 16 kusů ve 7 typech (z DXF přesných CAD bloků). 9 oken s žaluzií "
            "v purenitovém kastlíku (do ulice Fibichova), 7 oken bez žaluzií (do "
            "zahrady). Vstupní dveře: 2 kusy (ulice + zahrada). Vnitřní dveře: "
            "odhadem 15 kusů (3 byty × ~5 dveří). Klempířina: oplechování krytiny, "
            "parapety u 16 oken, dešťové svody, vikýřové prvky."
        ),
        "vliv": (
            "Položky 'PSV-76 Výplně otvorů' a 'PSV-76 Klempířina' v rozpočtu mají "
            "přesné počty u oken (DXF), ale dveře jsou odhadem ±20 %."
        ),
    },
    6: {
        "title": "Detail vytápění — kamna, kotel, krb, tepelné čerpadlo",
        "komu": "Investor Volný (rozhodnutí o konkrétních výrobcích a kapacitách)",
        "komu_email": "JindraVolny@seznam.cz",
        "vaznost": "DŮLEŽITÉ",
        "o_co_jde": (
            "Technická zpráva uvádí 4 zdroje vytápění:\n"
            "  • Sporáková kamna na tuhá paliva v 1.NP s akumulačním zásobníkem TUV\n"
            "  • Elektrokotel pro 3.NP (samostatný byt)\n"
            "  • Krb na tuhá paliva v 3.NP\n"
            "  • Multisplit tepelné čerpadlo vzduch-vzduch (venkovní + vnitřní jednotky)\n\n"
            "Ale technická zpráva neuvádí konkrétní výkony v kW, výrobce ani počet "
            "vnitřních jednotek tepelného čerpadla. Investorovo rozhodnutí potřebné "
            "pro objednání zařízení."
        ),
        "co_je_treba": (
            "Karle, prosím přepošli Volnému tyto otázky:\n"
            "  • Sporáková kamna: výrobce + výkon (např. 10 kW Atmos)?\n"
            "  • Elektrokotel pro 3.NP: jaký výkon (8–12 kW)?\n"
            "  • Krb: jaký výrobce a výkon (5–8 kW)?\n"
            "  • Multisplit TČ: kolik vnitřních jednotek (4 nebo 5)? Výrobce (Daikin/LG)?\n"
            "  • Akumulační zásobník TUV: jaký objem (300/500/750 l)?"
        ),
        "co_mame": (
            "Momentálně použito jako odhad podle běžných výrobků pro RD 220 m² "
            "se 3 bytovými jednotkami:\n"
            "  • Kamna ~10 kW, elektrokotel ~8 kW, krb ~6 kW\n"
            "  • Multisplit TČ: 1 venkovní + 5 vnitřních jednotek\n"
            "  • Akumulační zásobník TUV: 300 l v 1.PP + elektrický bojler 80 l v 3.NP"
        ),
        "vliv": (
            "Položky 'PSV-73 Vytápění' v rozpočtu mají odhadové ceny v rozsahu "
            "±30 % skutečné ceny. Po upřesnění Volnému se přesnost zvýší."
        ),
    },
    7: {
        "title": "Rozsah elektroinstalace — počty zásuvek, svítidel, příprava FVE",
        "komu": "Elektroprojektant + Investor Volný",
        "komu_email": "Karel → vybrat elektroprojektanta nebo Volnému k upřesnění",
        "vaznost": "DŮLEŽITÉ",
        "o_co_jde": (
            "Technická zpráva uvádí: „kompletně nová elektroinstalace v celém domě, "
            "nová pojistková skříň s podružným měřením pro každé patro, příprava na "
            "osazení FVE (fotovoltaiky)“. Ale konkrétní počty zásuvek, svítidel, "
            "okruhů, ani metráž kabeláže nejsou."
        ),
        "co_je_treba": (
            "Karle, doporučujeme jednu z těchto možností:\n"
            "1) Najít elektroprojektanta a zpracovat profesní část (D.1.4 Silnoproud) — "
            "    cena 30–50 tis. Kč, výsledkem přesný rozpočet ELI\n"
            "2) Použít odhad podle standardu pro RD 220 m² se 3 byty, s rezervou "
            "    v cenové nabídce ~15 % na variabilitu"
        ),
        "co_mame": (
            "Momentálně použito jako odhad:\n"
            "  • 70 zásuvek (60–80 podle rozsahu)\n"
            "  • 35 svítidel (30–40)\n"
            "  • 4 rozváděče (1 hlavní + 3 podružné pro patra)\n"
            "  • Rozvody silnoproudu cca 700 bm (RD 220 m² × 3,2 bm/m² standard)\n"
            "  • Příprava FVE: rezerva v rozváděči + trasy do střechy"
        ),
        "vliv": (
            "Položky 'M-21 Elektroinstalace' mají odhadové výměry s přesností ±20 %."
        ),
    },
    8: {
        "title": "Bilance výztuže ŽB konstrukcí — kg/m³ pro každý prvek",
        "komu": "Statik TeAnau (Ing. Jan Tvardík, ČKAIT 0012219)",
        "komu_email": "TeAnau s.r.o.",
        "vaznost": "STŘEDNĚ DŮLEŽITÉ",
        "o_co_jde": (
            "Statická TZ uvádí třídy materiálu (beton C25/30, výztuž B500B, ocel S235) "
            "a rozměry konstrukcí, ale nepiše konkrétní spotřebu výztuže v kg/m³ "
            "betonu pro jednotlivé prvky (opěrná stěna bílá vana, pozední věnec, "
            "deska na terénu, schodišťová deska)."
        ),
        "co_je_treba": (
            "Karle, prosíme TeAnau o krátkou tabulku spotřeby výztuže B500B "
            "pro každý ŽB prvek (kg/m³ betonu). Pokud má statik výztužné výkresy "
            "z RFEM výpočtu, ještě lépe — z nich vyplyne přesná hmotnost."
        ),
        "co_mame": (
            "Momentálně použito (empirické sazby Methvin pro tyto typy konstrukcí):\n"
            "  • Opěrná stěna bílá vana W0: ~120 kg/m³\n"
            "  • Pozední věnec: ~100 kg/m³\n"
            "  • Deska podlahy 1.NP: ~75 kg/m³"
        ),
        "vliv": (
            "Položky 'HSV-2 Železobeton — výztuž' mají odhadové kg. Po upřesnění "
            "od TeAnau se přesnost zvýší. Rozdíl ±30 % v kg = ±10 % v ceně výztuže."
        ),
    },
    9: {
        "title": "Bourací práce — přesné objemy klenby, podlah, krovu",
        "komu": "Informativní — Karel rozhodne sám podle praxe",
        "komu_email": "—",
        "vaznost": "STŘEDNĚ DŮLEŽITÉ",
        "o_co_jde": (
            "Souhrnná zpráva v kap. m.10.e uvádí celkové tonáže odpadu z bourání:\n"
            "  • Beton / cihly / keramika: 46 t\n"
            "  • Dřevo: 9 t\n"
            "  • Kovy: 0,1 t\n"
            "  • EPS: 0,1 t\n"
            "  • Sádra: 0,2 t\n"
            "  • Směsné stavební: 1,0 t\n"
            "Celkem 56,4 t.\n\n"
            "To je dobrá hodnota pro VRN (kontejnery, doprava na skládku), ale pro "
            "rozpočet bourací práce by ideálně bylo m³/m² každého prvku samostatně."
        ),
        "co_je_treba": (
            "Žádná akce směrem k projektantovi. Karle, použijeme empirický rozklad "
            "celkových tonáží na jednotlivé položky podle objemových hmotností. "
            "Pro zhotovitele zkušeného s rekonstrukcemi je toto dostačující."
        ),
        "co_mame": (
            "Momentálně použito rozklad tonáží:\n"
            "  • Bourání krovu: 9 t × 0,75 t/m³ dřeva = 12 m³\n"
            "  • Bourání střešní krytiny: ~141 m² plechu\n"
            "  • Bourání nadezdívek: ~6 m³ cihel\n"
            "  • Bourání trámového stropu 2.NP/podkroví: 104,4 m²\n"
            "  • Bourání podlah 1.NP+2.NP: ~154 m²\n"
            "  • Likvidace odpadu: 56,9 t celkem (VRN)"
        ),
        "vliv": "Žádný — empirický rozklad je standardní postup pro DSP rozsah.",
    },
    10: {
        "title": "Zařízení staveniště — umístění, zábory ulice Fibichova",
        "komu": "Karel sám (organizační — neblokuje rozpočet)",
        "komu_email": "—",
        "vaznost": "MÉNĚ URGENTNÍ",
        "o_co_jde": (
            "Souhrnná technická zpráva kap. m.10 uvádí: „staveniště bude zřízeno přímo "
            "v budově a na pozemcích investora“. Ale ulice Fibichova je úzká, řadová "
            "zástavba, bude třeba krátkodobé zábor pro dopravu materiálu."
        ),
        "co_je_treba": (
            "Karle, prosíme tě po podpisu smlouvy s investorem vyřídit krátkodobé "
            "zábory ulice Fibichova u Městského úřadu Jáchymov + dopravní inspektorátu. "
            "Předpoklad: cca 30 dnů kumulativně."
        ),
        "co_mame": (
            "Momentálně použito v rozpočtu VRN ~5–7 % z přímých nákladů (běžný rozsah "
            "pro RD rekonstrukci v městské zástavbě), z toho:\n"
            "  • Buňky kancelář + sociální: 8 měsíců\n"
            "  • Mobilní WC: 8 měsíců\n"
            "  • Krátkodobé zábory: paušál\n"
            "  • Kontejnery na suť (7 ks velkoobjemových)\n"
            "  • BOZP koordinace, pojištění CAR\n"
            "  • Doprava materiálu (cca 2–3 % z přímých nákladů)"
        ),
        "vliv": "Žádný — empirický koeficient je obvyklý pro tento typ stavby.",
    },
    11: {
        "title": "Lokální cenová hladina Karlovarský kraj",
        "komu": "Informativní — Karel/Aleksandr rozhodne při finalizaci",
        "komu_email": "—",
        "vaznost": "MÉNĚ URGENTNÍ",
        "o_co_jde": (
            "Jáchymov leží v Krušnohoří, sněhová oblast VII. (nejvyšší v ČR), "
            "obtížnější dostupnost pro dodavatele než např. Praha nebo Karlovy Vary. "
            "Lokální cenová hladina cca +5 až +10 % oproti pražskému průměru."
        ),
        "co_je_treba": (
            "Žádná akce projektantovi. Při finalizaci cenové nabídky aplikujeme buď:\n"
            "  • Regionální koeficient z ÚRS pro Karlovarský kraj (pokud dostupný)\n"
            "  • Plošně +5 % k republikovému průměru jako konzervativní odhad"
        ),
        "co_mame": "Cenovou hladinu finalizujeme po výběru rozsahu rozpočtu (otázka 1).",
        "vliv": "Pouze finální cenování, neovlivňuje výměry ani rozsah práce.",
    },
    12: {
        "title": "Status žádosti o stavební povolení",
        "komu": "Informativní — sledovat průběh, neblokuje rozpočet",
        "komu_email": "—",
        "vaznost": "INFORMATIVNÍ",
        "o_co_jde": (
            "Status k 10. 4. 2026:\n"
            "  • 260217 sklad + parking: podáno 17. 2. 2026, čeká se na rozhodnutí "
            "(po termínu, projektant urguje)\n"
            "  • 260219 dům: podáno 23. 3. 2026, čeká"
        ),
        "co_je_treba": (
            "Žádná akce ze strany rozpočtu. Rozpočet připravujeme pro realizaci "
            "2026–2027 podle technické zprávy. Po vydání stavebního povolení bude "
            "možné zahájit smluvní jednání s investorem."
        ),
        "co_mame": "Realizace plánována 2026–2027 dle souhrnné TZ kap. m.1.m.",
        "vliv": "Žádný — pouze časový plán.",
    },
    13: {
        "title": "Sklad + parkování — chybí řezy, pohledy, návrhové půdorysy",
        "komu": "Architekt SMASH (Ing. arch. Marek Smolka)",
        "komu_email": "SMASH architekti s.r.o.",
        "vaznost": "STŘEDNĚ DŮLEŽITÉ",
        "o_co_jde": (
            "Pro sklad + parkování máme pouze 1 architektonický výkres "
            "(D.1.1.02.R1 — Půdorys suterénu skladu). Chybí řezy, pohledy a návrhové "
            "půdorysy. Sklad je utilitární objekt (lichoběžník 6,35×3,34 m podle "
            "DXF) a může to být úmysl projektanta — považován za „jednoduchý "
            "technický prostor“. Ale chtěli bychom mít jistotu."
        ),
        "co_je_treba": (
            "Karle, prosíme tě ověřit u Bc. Jiřího Šmída nebo SMASHE, zda existuje "
            "další architektonická dokumentace skladu (řezy, pohledy). Pokud ano, "
            "rádi bychom ji obdrželi. Pokud ne, akceptujeme odhad ze statické TZ "
            "a DXF rozměrů."
        ),
        "co_mame": (
            "Momentálně použito: rozměry sklad 6,35 × 3,34 m (z DXF přesné), "
            "parkování 7,0 m délka × 7 IPE180 stojek (z DXF), obvod základu 19,4 m. "
            "Cca 27 položek v rozpočtu pro sklad — odhad převzat z TZ + DXF."
        ),
        "vliv": (
            "Sklad scope v rozpočtu má přesnost ±15 %. Pokud doplnit dokumentaci, "
            "přesnost vzroste na ±5 %."
        ),
    },
    14: {
        "title": "Statický výpočet D.2.2 — pouze pro dům, nebo i sklad?",
        "komu": "Statik TeAnau",
        "komu_email": "TeAnau s.r.o.",
        "vaznost": "STŘEDNĚ DŮLEŽITÉ",
        "o_co_jde": (
            "Po roztřídění dokumentace jsme zjistili, že soubor D.2.2 — Statický "
            "výpočet (19,8 MB plné znění s přílohami) v dokumentaci nemá objektový "
            "sufix (dum/sklad) v názvu. Sklad má samostatnou statickou TZ "
            "(D.2.1 Statika sklad), takže pravděpodobně D.2.2 je výpočet pouze "
            "pro dům. Ale rádi bychom ověřili."
        ),
        "co_je_treba": (
            "Karle, prosíme tě krátký dotaz u TeAnau (Ing. Tvardík): "
            "„Je statický výpočet D.2.2 (19,8 MB) zpracován pouze pro dům 260219, "
            "nebo zahrnuje i sklad 260217?“\n"
            "Pokud zahrnuje sklad, máme možnost vytěžit ještě více detailů."
        ),
        "co_mame": (
            "Momentálně použito: D.2.2 pokrývá pouze dům. Sklad statika je v menším "
            "samostatném výpočtu uvnitř D.2.1 Statika sklad."
        ),
        "vliv": (
            "Pouze upřesnění zdrojů. Pokud sklad nemá svůj výpočet, mohli bychom o "
            "to TeAnau požádat (cena 5–10 tis. Kč), abychom měli přesnější rozpočet."
        ),
    },
    15: {
        "title": "C.01 Situace širších vztahů — dvě paralelní verze názvu",
        "komu": "Informativní — Karel (drobnost, neblokuje)",
        "komu_email": "—",
        "vaznost": "MÉNĚ URGENTNÍ",
        "o_co_jde": (
            "Při třídění dokumentace jsme našli dvě paralelní verze stejné situace "
            "s mírně odlišnými názvy: „C.01 Situace širších vztahů“ (721 KB) a "
            "„C.01 Situační výkres širších vztahů“ (724 KB). Téměř identická "
            "velikost, liší se jen formálností názvu. Pravděpodobně to jsou dvě "
            "úvodní verze pojmenování, ne dvě paralelní revize."
        ),
        "co_je_treba": "Žádná akce. Použijeme tu s formálnějším názvem („Situační výkres“) jako finální.",
        "co_mame": (
            "V kanonické struktuře dokumentace ponecháno: „C.01 Situační výkres "
            "širších vztahů_EAR.pdf“ (formální název dle vyhlášky 499/2006 Sb. "
            "příloha č. 1). Druhá verze v archivu _superseded pro audit trail."
        ),
        "vliv": "Žádný — administrativní detail.",
    },
    16: {
        "title": "Interní: dům/sklad statika TZ byly v dokumentaci vyměněny — již opraveno",
        "komu": "Informativní pro Karla (interní záznam, neoznamovat projektantům)",
        "komu_email": "—",
        "vaznost": "INFORMATIVNÍ",
        "o_co_jde": (
            "Při nezávislé revizi technické dokumentace jsme zjistili, že soubor "
            "v kanonické cestě tz/260219_dum/D_2_1_TZ_statika_dum_TeAnau.pdf "
            "obsahoval ve skutečnosti obsah statiky pro sklad (hlavička každé "
            "stránky uváděla „Zahradní sklad...“), a opačně. Soubory byly "
            "v dokumentaci fyzicky prohozené. To je chyba ze zpracování dokumentace "
            "na naší straně (při uploadu zip souboru), nikoli chyba v originální PD "
            "od TeAnau.\n\n"
            "Po prohození souborů zpět do správných cest se přesnost ověření "
            "technických údajů zvýšila z 52 % (před opravou) na 97 % (po opravě). "
            "Bíla vana, HEA180/200, Porotherm 30 a ostatní detaily nyní správně "
            "mapovány na soubor o domě."
        ),
        "co_je_treba": (
            "Žádná akce. Pouze informativní záznam pro Karla, aby věděl, že jsme tuto "
            "data quality issue identifikovali a opravili před vygenerováním rozpočtu."
        ),
        "co_mame": "Soubory již prohozeny zpět do správných cest. Rozpočet generován z opravené dokumentace.",
        "vliv": "Žádný — chyba opravena před zahájením rozpočtu.",
    },
    17: {
        "title": "Sloupky krovu — terminologie „jakl“ / „jekl“ vs. RHS/SHS",
        "komu": "Statik TeAnau (drobná terminologická korekce)",
        "komu_email": "TeAnau s.r.o.",
        "vaznost": "MÉNĚ URGENTNÍ",
        "o_co_jde": (
            "Statika dům používá pro sloupky pod středovou vaznicí termín "
            "„sloupky z jeklu 100/4“ (statika §4) nebo „jakl 100/4“ (ARS). "
            "„Jakl“ / „jekl“ je německý výrobní název (Jäckl) pro uzavřený "
            "obdélníkový průřez. Český normativní termín je RHS nebo SHS profil "
            "podle ČSN EN 10219-2 (uzavřený profil za studena tvarovaný)."
        ),
        "co_je_treba": (
            "Karle, drobnost. Pro položku v rozpočtu „Sloupky pod vaznice“ "
            "použijeme normativní terminologii: RHS / SHS 100×100×4 mm S235JR. "
            "Funkčně shodné s „jakl 100/4“, jen formálně správně."
        ),
        "co_mame": (
            "V rozpočtu položka „HSV-5 Sloupky pod vaznice — uzavřený profil "
            "100×100×4 mm“ (6 kusů × ~2,5 m výška × 11,7 kg/m = 176 kg)."
        ),
        "vliv": "Žádný — pouze normativní terminologie pro položku v rozpočtu.",
    },
    18: {
        "title": "Sklad geometrie — řešeno z DXF výkresů",
        "komu": "Informativní (vyřešeno automaticky z DXF)",
        "komu_email": "—",
        "vaznost": "INFORMATIVNÍ",
        "o_co_jde": (
            "Technická zpráva uvádí rozměry skladu „lichoběžník 6,35 × 3,34 m“ a "
            "„parkovací stání délky 7,0 m“ pouze jako sumární údaje v textu, "
            "ne jako kótované rozměry. Pro přesný rozpočet jsme potřebovali tyto "
            "hodnoty ověřit."
        ),
        "co_je_treba": "Žádná akce — vyřešeno z DXF výkresů.",
        "co_mame": (
            "Všechny tři rozměry verifikovány z DXF výkresů s vysokou přesností:\n"
            "  • Sklad šířka 6,35 m = 6350,06 mm (DXF DIMENSION, sklad_DPZ)\n"
            "  • Sklad hloubka 3,34 m = exactly 3340,0 mm (DXF DIMENSION, dum_DPZ)\n"
            "  • Parkování délka 7,0 m = 7000,0 mm (DXF LWPOLYLINE bbox, 4 polylines "
            "    potvrzují stejnou hodnotu)"
        ),
        "vliv": (
            "Položka „HSV-4 IPE180 parking sekundární zastřešení“ v rozpočtu používá "
            "7 kusů × 7 m × 18,8 kg/m s confidence 0,90 (vysokou)."
        ),
    },
    19: {
        "title": "Strop S09 — „košický plech“ vs „trapéz 40S/160“ (terminologický rozpor)",
        "komu": "Statik TeAnau s.r.o. (Tvardík, Bendík)",
        "komu_email": "TeAnau s.r.o.",
        "vaznost": "STŘEDNĚ DŮLEŽITÉ",
        "o_co_jde": (
            "ARS řez S09 (skladba stropu 1.NP/2.NP) uvádí „košický plech“ jako konstrukci "
            "spřaženého stropu. TZ statika dům §6.3 ale uvádí přesnou specifikaci "
            "„trapézový plech 40S/160 tl. 0,75 mm + nabetonávka 60 mm z betonu C25/30 XC1“. "
            "Jde o stejnou konstrukci (jen synonymum), nebo o dvě různé varianty?"
        ),
        "co_je_treba": (
            "Karle, prosím dotaž TeAnau, zda je „košický plech“ a „trapéz 40S/160 + nabetonávka“ "
            "synonymum (= jedna a tatáž ocelobetonová deska s trny pro spřažení) nebo zda jde "
            "o dvě různé skladby. Pokud synonymum, sjednotit terminologii v dokumentaci."
        ),
        "co_mame": (
            "Pracovní hypotéza: synonyma. „Košický plech“ je hovorové označení pro trapézový "
            "plech s trny / vrubovými žebry pro spřažení s nabetonovanou monolitickou deskou — "
            "typická skladba ocelobetonového stropu v ČR. V rozpočtu položeno jako "
            "„HSV-4 Ocelobetonový strop — trapéz 40S/160 tl. 0,75 mm + nabetonávka 60 mm C25/30“ "
            "podle specifické varianty z TZ statika (preferujeme přesnější popis)."
        ),
        "vliv": (
            "Pokud synonymum (předpokládaná varianta): položka HSV-4 ocelobetonový strop "
            "je správná. Pokud dvě skladby: zhotovitel musí ověřit, která je závazná — "
            "ARS nebo TZ statika. Standardně preferuje TZ statika."
        ),
    },
    20: {
        "title": "Verifikace ÚRS kódů — production lookup před cenotvorbou",
        "komu": "Investor Volný + Karel (rozhodnutí o produkčním ověření)",
        "komu_email": "JindraVolny@seznam.cz",
        "vaznost": "STŘEDNĚ DŮLEŽITÉ",
        "o_co_jde": (
            "Rozpočet obsahuje 189 položek s navrženými ÚRS kódy generovanými heuristikou "
            "v Phase 1. V Part 5b (Sheet 9 Var_F_URS_Verification_Trail) jsme ověřili "
            "12 kódů přes WebSearch (veřejná zrcadla podminky.urs.cz: smlouvy.gov.cz, "
            "vhodne-uverejneni.cz, docplayer.cz, cs-urs.cz).\n\n"
            "Výsledek ověření: heuristika odhadne první 6 cifer (rodinu) správně v ~75 % "
            "případů, ale 9-cifrový leaf je chybný v ~63 % (odlišný distance band, "
            "materiál, geometrie, nebo lokace). 3 položky verifikovány jako přesné, "
            "9 položek označeno LEAF CHYBNÝ (viz červeně označený sloupec v Sheet 5/6 "
            "Var_B_Polozkovy), 4 nahrazeny správnými kódy (HSV6.007 keramické obklady → "
            "781473810 atd.)."
        ),
        "co_je_treba": (
            "Pro produkční cenotvorbu doporučujeme spustit zbývajících 177 položek "
            "přes URS_MATCHER service (Perplexity-driven online lookup s ÚRS katalogem 2026/I). "
            "Náklad: ~3-4 hodiny processing + 30-60 min cleanup ≈ 3-4 tis. Kč extra service.\n\n"
            "Karle, prosím přepošli Volnému s žádostí o rozhodnutí: ANO objednat full URS "
            "verification před cenovou nabídkou, NEBO ponechat na zhotoviteli s upozorněním "
            "v cenové nabídce, že kódy jsou heuristické a vyžadují ověření."
        ),
        "co_mame": (
            "Aktuální stav (po Part 5b WebSearch):\n"
            "  • 4 položky urs_status=matched_websearch_verified ✓ (confidence 0,90–0,95)\n"
            "  • 9 položek urs_status=wrong_leaf_disambiguation_needed ⚠ (confidence snížena "
            "na 0,50, 6-ciferná rodina v urs_code_family_6digit)\n"
            "  • 176 položek urs_status=needs_production_lookup (confidence 0,65, čeká "
            "na URS_MATCHER service run)\n\n"
            "Excel Sheet 9 Var_F obsahuje plný audit trail per kód: verdict (matched / "
            "wrong leaf / correct replacement), verified URS popis, item popis, "
            "correct_code_hint pro disambiguation."
        ),
        "vliv": (
            "Pokud URS verification objednán: cenová nabídka má production-grade ÚRS kódy "
            "(~95 % top-1 acceptance), zhotovitel ji může bez úprav předložit. "
            "Pokud NE: cena platí, ale zhotovitel musí v provádění nahradit ~120 chybných "
            "leaves vlastní inženýrskou cestou — rezerva 5-8 % v ceně doporučena."
        ),
    },
}


def set_cell_bg(cell, color_hex: str):
    """Set table cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def vaznost_color(vaznost: str) -> str:
    """Map vaznost label to hex color for table cell."""
    return {
        "URGENTNÍ":          "FFCDD2",   # light red
        "DŮLEŽITÉ":          "FFE0B2",   # light orange
        "STŘEDNĚ DŮLEŽITÉ":  "FFF9C4",   # light yellow
        "MÉNĚ URGENTNÍ":     "C8E6C9",   # light green
        "INFORMATIVNÍ":      "E0E0E0",   # light gray
    }.get(vaznost, "FFFFFF")


def add_page_number(paragraph):
    """Append a page-number field code to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def main() -> int:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("RD Jáchymov Fibichova 733")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Otázky pro vyřešení s projektanty před zahájením realizace")
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # Header info table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Light Grid Accent 1"
    info_data = [
        ("Investor", "Mgr. Jindřich Volný"),
        ("Zhotovitel", "Karel Šmíd, smid.karell@gmail.com, +420 608 930 914"),
        ("Rozpočtář", "Aleksandr Pro (STAVAGENT)"),
        ("Datum", "18. 5. 2026"),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        info_table.cell(i, 1).text = value
    for row in info_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.add_run(
        "Stavba je rozdělena do dvou stavebních akcí:\n"
        "  • 260219 Řadový rodinný dům Fibichova 733 — rekonstrukce + nástavba 3.NP\n"
        "  • 260217 Zahradní sklad + parkovací stání"
    )

    intro2 = doc.add_paragraph()
    intro2.add_run(
        "Tento dokument obsahuje 18 otázek, které je třeba vyřešit s projektanty "
        "před cenovou kalkulací konkrétními řemesly a před zahájením realizace.\n\n"
        "Každá otázka má uvedeno: kdo má odpovědět, co je třeba zjistit, jaký předpoklad "
        "jsme prozatím použili pro výpočet, a jak odpověď ovlivní rozpočet."
    )

    doc.add_paragraph()
    legend = doc.add_paragraph()
    legend.add_run("Legenda důležitosti:").bold = True
    legend_table = doc.add_table(rows=1, cols=5)
    legend_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, label in enumerate(["URGENTNÍ", "DŮLEŽITÉ", "STŘEDNĚ DŮLEŽITÉ", "MÉNĚ URGENTNÍ", "INFORMATIVNÍ"]):
        cell = legend_table.cell(0, i)
        cell.text = label
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, vaznost_color(label))

    doc.add_page_break()

    # Souhrnná tabulka
    doc.add_heading("Souhrnná tabulka otázek", level=1)
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = "Light Grid Accent 1"

    hdr = summary_table.rows[0].cells
    for i, h in enumerate(["č.", "Kdo má odpovědět", "Krátké shrnutí", "Důležitost"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], "1F3A5F")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for qid in sorted(OTAZKY):
        q = OTAZKY[qid]
        row = summary_table.add_row().cells
        row[0].text = str(qid)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = q["komu"]
        row[2].text = q["title"]
        row[3].text = q["vaznost"]
        set_cell_bg(row[3], vaznost_color(q["vaznost"]))
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # Column widths
    summary_table.columns[0].width = Cm(1.0)
    summary_table.columns[1].width = Cm(5.5)
    summary_table.columns[2].width = Cm(7.5)
    summary_table.columns[3].width = Cm(3.0)

    doc.add_page_break()

    # Detail per otázka
    doc.add_heading("Detailní popis otázek", level=1)
    for qid in sorted(OTAZKY):
        q = OTAZKY[qid]
        # Heading
        h = doc.add_heading(f"Otázka č. {qid}:  {q['title']}", level=2)

        # Metadata table
        meta = doc.add_table(rows=2, cols=2)
        meta.style = "Light List Accent 1"
        meta.cell(0, 0).text = "Komu"
        meta.cell(0, 0).paragraphs[0].runs[0].bold = True
        meta.cell(0, 1).text = q["komu"]
        meta.cell(1, 0).text = "Důležitost"
        meta.cell(1, 0).paragraphs[0].runs[0].bold = True
        meta.cell(1, 1).text = q["vaznost"]
        set_cell_bg(meta.cell(1, 1), vaznost_color(q["vaznost"]))
        meta.cell(1, 1).paragraphs[0].runs[0].bold = True

        # Status banner — for RESOLVED / PARTIALLY_RESOLVED questions
        if q.get("status"):
            p = doc.add_paragraph()
            status_label = q["status"].replace("_", " ")
            r = p.add_run(f"✓ STATUS: {status_label}")
            r.bold = True
            r.font.color.rgb = RGBColor(0x0F, 0x51, 0x32) if q["status"] == "RESOLVED" else RGBColor(0xA0, 0x40, 0x00)
            if q.get("status_note"):
                doc.add_paragraph(q["status_note"])

        # O co jde
        p = doc.add_paragraph()
        p.add_run("O co jde:").bold = True
        doc.add_paragraph(q["o_co_jde"])

        # Co je třeba
        p = doc.add_paragraph()
        p.add_run("Co je třeba:").bold = True
        doc.add_paragraph(q["co_je_treba"])

        # Co prozatím máme
        p = doc.add_paragraph()
        p.add_run("Co prozatím máme (předpoklad pro výpočet):").bold = True
        doc.add_paragraph(q["co_mame"])

        # Vliv
        p = doc.add_paragraph()
        p.add_run("Vliv na rozpočet:").bold = True
        doc.add_paragraph(q["vliv"])

        doc.add_paragraph()

    doc.add_page_break()

    # Závěrečná sekce — poznámky pro Karla
    doc.add_heading("Poznámky pro Karla — komu které otázky", level=1)

    karel_notes = [
        ("Pro investora pana Volného", "1 (rozsah rozpočtu), 2 (parkování), 6 (vytápění detail), 7 (rozsah ELI)"),
        ("Pro architekta SMASH (Ing. Smolka)", "3 (DPS doplnit), 4 (skladby podlah), 5 (výpisy oken/dveří/klempířiny), 13 (sklad dokumentace)"),
        ("Pro statika TeAnau (Ing. Tvardík)", "8 (výztuž ŽB kg/m³), 14 (rozsah statického výpočtu D.2.2), 17 (terminologie sloupků krovu)"),
        ("Pro elektroprojektanta (vybrat)", "7 (rozsah ELI, počty zásuvek a svítidel)"),
        ("Karel sám organizačně řeší (neblokuje rozpočet)", "9 (bourací objemy), 10 (zařízení staveniště, zábory), 11 (cenová hladina)"),
        ("Pouze informativní — žádná akce", "12 (status SU), 15 (C.01 dvojice názvů), 16 (vnitřní záznam k souborům), 18 (sklad geometrie vyřešena z DXF)"),
    ]
    notes_table = doc.add_table(rows=1, cols=2)
    notes_table.style = "Light Grid Accent 1"
    hdr = notes_table.rows[0].cells
    hdr[0].text = "Komu adresovat"
    hdr[1].text = "Otázky č."
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
        set_cell_bg(c, "1F3A5F")
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for label, ids in karel_notes:
        row = notes_table.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = ids

    doc.add_paragraph()

    closing = doc.add_paragraph()
    closing.add_run(
        "Po doručení odpovědí od projektantů prosím o forward Alexandrovi (STAVAGENT) "
        "k aktualizaci rozpočtu. Každá vyřešená otázka zvýší přesnost příslušných "
        "položek v rozpočtu. Pro otázky 1, 2, 3 doporučujeme prioritní řešení — bez "
        "rozhodnutí Volného o rozsahu rozpočtu (otázka 1) nelze cenovou nabídku "
        "finalizovat.\n\n"
        "Děkuji za součinnost.\n\n"
        "Aleksandr Pro\n"
        "STAVAGENT — automated rozpočet pipeline"
    ).italic = True

    # Footer with page number + project name
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("RD Jáchymov Fibichova 733 — Otázky pro projektanty — 18. 5. 2026 — strana ").font.size = Pt(8)
    add_page_number(footer_p)
    footer_p.add_run(" — STAVAGENT auto-generated").font.size = Pt(8)

    # Save
    doc.save(str(OUT))
    print(f"✓ Wrote {OUT.relative_to(PROJ)} ({OUT.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    main()
