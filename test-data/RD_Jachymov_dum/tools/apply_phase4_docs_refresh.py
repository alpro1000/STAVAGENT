#!/usr/bin/env python3
"""
Phase 4 — VALIDATE LIST docs refresh.

Applies gate-3 D.5 + D.6 dispositions plus Q21 (plot dřevěný out-of-scope)
across MD outputs + Word otázky doc.

Surgical updates to preserve existing prose:

1. MD standardisation (D.5):
   - "208 položek total"  → "211 položek total (207 active, 4 deprecated audit-trail)"
   - "208 (204 active ...)" → "211 (207 active, 4 deprecated audit-trail)"
   - "208 items" → "211 items"
   - "204 active" → "207 active"
   - dum count "177" updates in OnePager + Summary to "180"
   - applies to 4 MDs: OnePager + Project_Summary + items_completeness_v2 + items_quality_report

2. Word otázky doc (D.6 + Q21):
   - intro "Tento dokument obsahuje 18 otázek" → "20 otázek (4 fully resolved:
     Q2, Q4, Q18, Q20 + 1 partially: Q5)"
   - Q2 summary-table status: "DŮLEŽITÉ" → "✅ RESOLVED — parking součást 260217"
   - Q20 summary-table status: "STŘEDNĚ DŮLEŽITÉ" → "✅ RESOLVED — Q20 závěr přijat"
   - Add Q21 row to summary table: "Plot dřevěný (133 INSERTs v DXF) — v scope?"
   - Add Q21 detail block at end of document

3. Project_Summary Q-status table (markdown):
   - Q2 / Q20 status → ✅ RESOLVED
   - Add Q21 row (plot dřevěný)
   - Final tally update

Idempotent: re-run is safe; replaced phrases are searched for both old and new
forms before applying.
"""

from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from copy import deepcopy

ROOT = Path(__file__).resolve().parent.parent
OUTPUTS = ROOT / "outputs"
TODAY = str(date.today())

# ---------------------------------------------------------------------------
# MD standardisation
# ---------------------------------------------------------------------------

# Order matters: more-specific patterns first so generic '208' doesn't pre-empt
# compound phrases.
MD_REPLACEMENTS: list[tuple[str, str]] = [
    # OnePager + Summary headline counts
    ("208 položek total", "211 položek total (207 active, 4 deprecated audit-trail)"),
    ("208 (204 active po audit v2)", "211 (207 active, 4 deprecated audit-trail)"),
    ("208 items — 177 dum + 27 sklad + 4 deprecated audit-trail", "211 items — 180 dum + 27 sklad + 4 deprecated audit-trail"),
    ("208 (204 active)", "211 (207 active, 4 deprecated audit-trail)"),
    ("208 items, 204 active", "211 items, 207 active"),
    ("177 dum + 27 sklad", "180 dum + 27 sklad"),
    ("SO 260219 — Dům** (177 položek)", "SO 260219 — Dům** (180 active položek)"),
    # Cell value "| **177** | **27** | **204** |"
    ("| **177** | **27** | **204** |", "| **180** | **27** | **207** |"),
    # Total-line table cells with counter
    ("Items celkem | 208 (204 active)", "Items celkem | 211 (207 active, 4 deprecated audit-trail)"),
    # Generic bare 208 → 211 (apply last, only when clearly an item-count context)
    ("208 položek", "211 položek"),
    ("208 items", "211 items"),
    ("(208 items)", "(211 items)"),
    # 204 active → 207 active (only when next to "204" item-count context)
    ("204 active po audit v2", "207 active (post audit v2 + per-drawing audit)"),
    # Resolved tally update
    ("**Resolved 2 (Q4 + Q18), partially resolved 1 (Q5), open 17.**",
     "**Resolved 4 (Q2 + Q4 + Q18 + Q20), partially resolved 1 (Q5), open 15 + 1 nový (Q21 plot dřevěný — out-of-scope working assumption per gate-2).**"),
    ("17 otevřených otázek (3 RESOLVED)",
     "15 otevřených otázek + 1 nový Q21 (4 RESOLVED + 1 partial)"),
]

# Q-status table cell updates in Project_Summary (markdown row format)
SUMMARY_TABLE_REPLACEMENTS: list[tuple[str, str]] = [
    (
        "| 2 | Parkovací stání u Dvořákovy — kam patří? | Investor + architekt | DŮLEŽITÉ |",
        "| 2 | Parkovací stání u Dvořákovy — kam patří? | Investor + architekt | ✅ **RESOLVED** — parking součást 260217 (assumption confirmed per gate-2) |",
    ),
    (
        "| 20 | Verifikace ÚRS kódů — production lookup | Investor + zhotovitel | STŘEDNĚ DŮLEŽITÉ |",
        "| 20 | Verifikace ÚRS kódů — production lookup | Investor + zhotovitel | ✅ **RESOLVED** — Q20 závěr přijat |",
    ),
]

# Q21 markdown row to insert after Q20 in Project_Summary table
Q21_MD_ROW = (
    "| 21 | Plot dřevěný (133 INSERTs v DXF) — v scope, nebo zahrada-only? | Investor "
    "(za Karla) | ℹ️ **WORKING ASSUMPTION: out-of-scope** (per CEV gate-2 disposition) |"
)

MD_TARGETS = [
    "Project_OnePager_RD_Jachymov.md",
    "Project_Summary_RD_Jachymov.md",
    "items_completeness_report_v2.md",
    "items_quality_report.md",
]


def patch_md(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    original = text
    changes: list[str] = []
    for old, new in MD_REPLACEMENTS:
        if old in text:
            text = text.replace(old, new)
            changes.append(old[:80])
    # Only apply to Project_Summary: Q-status table cell rewrites + Q21 row
    if path.name == "Project_Summary_RD_Jachymov.md":
        for old, new in SUMMARY_TABLE_REPLACEMENTS:
            if old in text:
                text = text.replace(old, new)
                changes.append(old[:80])
        # Insert Q21 row after Q20 row — only if not present already
        q20_anchor = "| 20 | Verifikace ÚRS kódů"
        if q20_anchor in text and "21 | Plot dřevěný" not in text:
            # Insert Q21 right after the end of the Q20 line
            lines = text.split("\n")
            new_lines: list[str] = []
            for ln in lines:
                new_lines.append(ln)
                if ln.startswith("| 20 | Verifikace ÚRS kódů"):
                    new_lines.append(Q21_MD_ROW)
                    changes.append("inserted Q21 row")
            text = "\n".join(new_lines)
    if text != original:
        path.write_text(text, encoding="utf-8")
    return {"file": path.name, "changes": changes, "applied": len(changes) > 0}


# ---------------------------------------------------------------------------
# Word otázky docx update
# ---------------------------------------------------------------------------

OTAZKY_DOCX_PATH = OUTPUTS / "Otazky_pro_Karla_a_projektanty_2026-05-18.docx"
OTAZKY_DOCX_V2_PATH = OUTPUTS / "Otazky_pro_Karla_a_projektanty_2026-05-26.docx"

INTRO_OLD = "Tento dokument obsahuje 18 otázek"
INTRO_NEW = (
    "Tento dokument obsahuje 20 otázek (4 plně vyřešené: Q2, Q4, Q18, Q20 + "
    "1 částečně vyřešená: Q5) + 1 nová Q21 (plot dřevěný — výchozí pracovní "
    "předpoklad: mimo scope)"
)

Q21_TITLE = "Otázka č. 21:  Plot dřevěný — v scope, nebo zahrada-only?"
Q21_O_CO_JDE = (
    "Per CEV per-drawing audit (commit 91ab8d2) jsme zjistili, že v DXF dum_DPZ "
    "+ sklad_DPZ je 133 INSERT bloků kategorie plot_dreveny. Není jasné, zda "
    "tyto dřevěné ploty patří do scope rozpočtu (Karel je staví) nebo jsou "
    "zahrada-only (investor řeší samostatně, nebo se jen vyznačují existující stav)."
)
Q21_CO_JE_TREBA = (
    "Karle, prosím ověř u SMASH nebo Volného. Pokud jsou součástí scope, "
    "doplníme HSV-X plot dřevěný item (~133 ks × průměrná délka × cena/m). "
    "Pokud zahrada-only, žádná akce."
)
Q21_CO_MAME = (
    "Pracovní předpoklad (per gate-2 disposition Alexandry): plot dřevěný = "
    "mimo scope, items.json nepřidáno. Pokud potvrzení jiné, rozpočet rozšíříme."
)
Q21_VLIV = (
    "Středně důležité. Pokud out-of-scope, žádný impact. Pokud in-scope, "
    "přidání +N položek HSV-X (rough estimate ~50-100 tis. Kč podle průměrné "
    "délky plotu)."
)


def _add_summary_row(table, q_no: int, popis: str, komu: str, status: str) -> None:
    new_row = table.add_row()
    cells = new_row.cells
    if len(cells) >= 4:
        cells[0].text = str(q_no)
        cells[1].text = popis
        cells[2].text = komu
        cells[3].text = status


def _find_paragraph_index(doc: Document, predicate) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text):
            return i
    return None


def patch_word_docx() -> dict:
    if not OTAZKY_DOCX_PATH.exists():
        return {"docx_path": str(OTAZKY_DOCX_PATH), "error": "source docx not found"}

    doc = Document(str(OTAZKY_DOCX_PATH))
    notes: list[str] = []

    # 1) Intro paragraph replacement
    for p in doc.paragraphs:
        if INTRO_OLD in p.text:
            # Replace run-by-run if possible (preserve formatting); fallback to
            # paragraph-level text rewrite.
            full = p.text
            new_full = full.replace(INTRO_OLD, INTRO_NEW)
            # Clear existing runs and set new text on a single run
            for r in list(p.runs):
                r.text = ""
            p.add_run(new_full)
            notes.append("intro paragraph rewritten")
            break

    # 2) Summary table — Otazky_pro_Karla docx has Table 2 with 4 columns
    #    ['č.', 'Kdo má odpovědět', 'Krátké shrnutí', 'Důležitost'] and 21 rows
    #    (header + Q1..Q20). Detection: 4-col table where row[1].cells[0]=="1"
    #    and row[1].cells[2] contains "rozpočt" (Q1 popis).
    summary_table = None
    for t in doc.tables:
        if len(t.rows) < 2 or len(t.rows[0].cells) != 4:
            continue
        row1 = t.rows[1].cells
        if row1[0].text.strip() == "1" and "rozpočt" in row1[2].text.lower():
            summary_table = t
            break

    if summary_table is None:
        notes.append("summary table not found")
    else:
        # Update Q2 + Q20 status cells
        for r in summary_table.rows[1:]:
            cells = r.cells
            if len(cells) < 4:
                continue
            qno = cells[0].text.strip()
            if qno == "2":
                cells[3].text = "✅ RESOLVED — parking součást 260217 (per gate-2)"
                notes.append("Q2 status updated")
            elif qno == "20":
                cells[3].text = "✅ RESOLVED — Q20 závěr přijat"
                notes.append("Q20 status updated")
        # Append Q21 row if not already present
        existing_q21 = any(r.cells[0].text.strip() == "21" for r in summary_table.rows[1:] if len(r.cells) >= 1)
        if not existing_q21:
            _add_summary_row(
                summary_table,
                21,
                "Plot dřevěný (133 INSERTs v DXF) — v scope?",
                "Investor (za Karla)",
                "ℹ️ WORKING ASSUMPTION: out-of-scope (per gate-2)",
            )
            notes.append("Q21 summary row added")

    # 3) Append Q21 detail block at end of document
    has_q21_block = any("Otázka č. 21" in p.text for p in doc.paragraphs)
    if not has_q21_block:
        doc.add_paragraph()  # spacer
        h = doc.add_paragraph()
        run = h.add_run(Q21_TITLE)
        run.bold = True
        run.font.size = Pt(13)
        doc.add_paragraph("O co jde:").runs[0].bold = True
        doc.add_paragraph(Q21_O_CO_JDE)
        doc.add_paragraph("Co je třeba:").runs[0].bold = True
        doc.add_paragraph(Q21_CO_JE_TREBA)
        doc.add_paragraph("Co prozatím máme (předpoklad pro výpočet):").runs[0].bold = True
        doc.add_paragraph(Q21_CO_MAME)
        doc.add_paragraph("Vliv na rozpočet:").runs[0].bold = True
        doc.add_paragraph(Q21_VLIV)
        notes.append("Q21 detail block appended")

    doc.save(str(OTAZKY_DOCX_V2_PATH))
    return {"docx_source": str(OTAZKY_DOCX_PATH.relative_to(ROOT)),
            "docx_v2_path": str(OTAZKY_DOCX_V2_PATH.relative_to(ROOT)),
            "notes": notes}


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------


def main() -> None:
    md_results: list[dict] = []
    for fname in MD_TARGETS:
        p = OUTPUTS / fname
        if not p.exists():
            md_results.append({"file": fname, "error": "missing"})
            continue
        md_results.append(patch_md(p))

    docx_result = patch_word_docx()

    log = {
        "applied_at": TODAY,
        "md_results": md_results,
        "docx_result": docx_result,
        "summary": {
            "mds_changed": sum(1 for m in md_results if m.get("applied")),
            "mds_unchanged": sum(1 for m in md_results if not m.get("applied", False) and "error" not in m),
            "mds_error": sum(1 for m in md_results if "error" in m),
        },
    }
    (OUTPUTS / "_phase4_docs_refresh_log.json").write_text(json.dumps(log, indent=2, ensure_ascii=False))
    print(json.dumps(log, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
