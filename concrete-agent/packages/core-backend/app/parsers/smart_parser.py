"""
Smart Parser - автоматический выбор оптимального парсера
Waterfall: pdfplumber -> MinerU (async) -> streaming

Логика:
 - PDF < 20MB -> pdfplumber (быстро)
 - PDF без текста (скан) -> MinerU async worker
 - Файл > 20MB -> Streaming парсеры (memory-efficient)
 - Автовыбор формата (Excel, PDF, XML)
"""
import asyncio
import logging
import os
import re
import shutil
import tempfile
import unicodedata
from pathlib import Path
from typing import Dict, Any, Optional

from app.parsers.excel_parser import ExcelParser
from app.parsers.pdf_parser import PDFParser
from app.parsers.kros_parser import KROSParser
from app.parsers.memory_efficient import (
    MemoryEfficientExcelParser,
    MemoryEfficientPDFParser,
    MemoryEfficientXMLParser
)

logger = logging.getLogger(__name__)

SIZE_THRESHOLD_MB = 20
ENABLE_MINERU = os.getenv("ENABLE_MINERU", "false").lower() == "true"


def _slugify(text: str) -> str:
    """
    Convert filename stem to ASCII-safe slug.
    Prevents MinerU crash on Windows with diacritics.
    e.g. 'IV MM-Ceník2026' -> 'IV_MM-Cenik2026'
    """
    normalized = unicodedata.normalize("NFKD", text)
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    safe = re.sub(r"[^\w\-]", "_", ascii_text)
    return safe.strip("_") or "document"


def _get_file_size_mb(file_path: Path) -> float:
    return file_path.stat().st_size / (1024 * 1024)


class SmartParser:

    def __init__(self):
        self.excel_parser = ExcelParser()
        self.pdf_parser = PDFParser()
        self.kros_parser = KROSParser()
        self.memory_excel = MemoryEfficientExcelParser()
        self.memory_pdf = MemoryEfficientPDFParser()
        self.memory_xml = MemoryEfficientXMLParser()

    def parse(self, file_path: str, file_type: Optional[str] = None, project_id: Optional[str] = None) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        detected_type = file_type or self._detect_type(path)
        if detected_type == "excel":
            return self.parse_excel(path)
        elif detected_type == "xml":
            return self.parse_xml(path)
        elif detected_type == "docx":
            return self.parse_docx(path)
        elif detected_type == "csv":
            return self.parse_csv(path)
        elif detected_type == "image":
            return self.parse_image(path)
        else:
            return self.parse_pdf(path)

    def parse_excel(self, path: Path) -> Dict[str, Any]:
        size_mb = _get_file_size_mb(path)
        logger.info(f"SmartParser: Excel {path.name} ({size_mb:.1f}MB)")
        if size_mb > SIZE_THRESHOLD_MB:
            return self.memory_excel.parse(str(path))
        return self.excel_parser.parse(str(path))

    def parse_xml(self, path: Path) -> Dict[str, Any]:
        size_mb = _get_file_size_mb(path)
        logger.info(f"SmartParser: XML {path.name} ({size_mb:.1f}MB)")
        if size_mb > SIZE_THRESHOLD_MB:
            return self.memory_xml.parse(str(path))
        return self.kros_parser.parse(str(path))

    def parse_docx(self, path: Path) -> Dict[str, Any]:
        """Parse DOCX using python-docx → text + tables."""
        logger.info(f"SmartParser: DOCX {path.name}")
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("python-docx not installed, falling back to PDF path")
            return {"positions": [], "text": "", "strategy": "docx_fallback"}

        try:
            doc = DocxDocument(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)

            tables_data: list = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                tables_data.append(rows)

            logger.info(f"SmartParser: DOCX extracted {len(paragraphs)} paragraphs, {len(tables_data)} tables")
            return {
                "positions": [],
                "text": full_text,
                "tables": tables_data,
                "strategy": "docx",
                "paragraphs": len(paragraphs),
            }
        except Exception as e:
            logger.error(f"SmartParser: DOCX failed: {e}")
            return {"positions": [], "text": "", "strategy": "docx_error", "error": str(e)}

    def parse_csv(self, path: Path) -> Dict[str, Any]:
        """Parse CSV/TSV → text + table."""
        import csv
        logger.info(f"SmartParser: CSV {path.name}")
        try:
            with open(str(path), "r", encoding="utf-8", errors="replace") as f:
                sample = f.read(4096)
                f.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
                except csv.Error:
                    dialect = csv.excel
                reader = csv.reader(f, dialect)
                rows = [row for row in reader]

            text_lines = [" | ".join(c for c in row if c.strip()) for row in rows if any(c.strip() for c in row)]
            logger.info(f"SmartParser: CSV extracted {len(rows)} rows")
            return {
                "positions": [],
                "text": "\n".join(text_lines),
                "tables": [rows] if rows else [],
                "strategy": "csv",
            }
        except Exception as e:
            logger.error(f"SmartParser: CSV failed: {e}")
            return {"positions": [], "text": "", "strategy": "csv_error", "error": str(e)}

    def parse_image(self, path: Path) -> Dict[str, Any]:
        """
        Parse image (JPG/PNG/TIFF) by converting to PDF, then using PDF pipeline.

        Flow: Pillow → save as single-page PDF → parse_pdf() pipeline
        This leverages pdfplumber → MinerU → memory_pdf fallback chain.
        """
        logger.info(f"SmartParser: Image {path.name}")
        try:
            from PIL import Image
            img = Image.open(str(path))
            # Convert to RGB if needed (RGBA/P modes can't save to PDF directly)
            if img.mode in ("RGBA", "P", "LA"):
                img = img.convert("RGB")
            elif img.mode not in ("RGB", "L"):
                img = img.convert("RGB")

            # Save as temporary PDF
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp_pdf = Path(tmp.name)
                img.save(str(tmp_pdf), "PDF", resolution=200.0)
                img.close()

            logger.info(f"SmartParser: Image → PDF ({tmp_pdf.stat().st_size / 1024:.0f}KB)")

            # Use the existing PDF pipeline (pdfplumber → MinerU → memory_pdf)
            result = self.parse_pdf(tmp_pdf)
            result["strategy"] = f"image_via_pdf ({result.get('strategy', 'unknown')})"
            result["source_image"] = path.name

            # Cleanup temp PDF
            try:
                tmp_pdf.unlink()
            except OSError:
                pass

            return result

        except ImportError:
            logger.warning("Pillow not installed, cannot parse images")
            return {"positions": [], "text": "", "strategy": "image_no_pillow"}
        except Exception as e:
            logger.error(f"SmartParser: Image failed: {e}")
            return {"positions": [], "text": "", "strategy": "image_error", "error": str(e)}

    def parse_pdf(self, path: Path) -> Dict[str, Any]:
        """
        Sync PDF parsing — pdfplumber → MinerU HTTP → memory_pdf fallback.

        MinerU is called via HTTP to mineru-service (separate Cloud Run)
        if MINERU_SERVICE_URL is set. Falls back to pdfplumber if not.
        """
        size_mb = _get_file_size_mb(path)
        logger.info(f"SmartParser: PDF {path.name} ({size_mb:.1f}MB)")
        if size_mb > SIZE_THRESHOLD_MB:
            return self.memory_pdf.parse(str(path))

        # Try pdfplumber first (fastest, local)
        try:
            result = self.pdf_parser.parse(str(path))
            positions = result if isinstance(result, list) else result.get("positions", [])
            if positions:
                logger.info(f"SmartParser: pdfplumber extracted {len(positions)} positions")
                return {"positions": positions, "strategy": "pdfplumber"}
            logger.info("SmartParser: pdfplumber returned 0 positions, trying MinerU...")
        except Exception as e:
            logger.warning(f"SmartParser: pdfplumber failed: {e}, trying MinerU...")

        # Try MinerU HTTP service (for scanned/hybrid PDFs)
        try:
            from app.parsers.mineru_client import parse_pdf_with_mineru
            mineru_text = parse_pdf_with_mineru(str(path), method="auto")
            if mineru_text and len(mineru_text) > 200:
                logger.info(f"SmartParser: MinerU extracted {len(mineru_text)} chars")
                return {
                    "text": mineru_text,
                    "positions": [],
                    "strategy": "mineru_service",
                    "text_length": len(mineru_text),
                }
        except Exception as e:
            logger.warning(f"SmartParser: MinerU service failed: {e}")

        # Final fallback
        return self.memory_pdf.parse(str(path))

    async def _parse_with_mineru_async(self, pdf_path: Path) -> Dict[str, Any]:
        safe_stem = _slugify(pdf_path.stem)
        output_dir = Path(tempfile.gettempdir()) / "mineru_output" / safe_stem
        output_dir.mkdir(parents=True, exist_ok=True)
        if safe_stem != pdf_path.stem:
            safe_pdf = output_dir / f"{safe_stem}.pdf"
            shutil.copy2(pdf_path, safe_pdf)
            source_path = safe_pdf
        else:
            source_path = pdf_path
        # MinerU 2.x: omit -b to use default hybrid-auto-engine backend
        cmd = ["mineru", "-p", str(source_path), "-o", str(output_dir), "-d", "cpu"]
        logger.info(f"SmartParser: running MinerU async: {' '.join(cmd)}")
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await proc.communicate()
        if proc.returncode != 0:
            err = stderr.decode("utf-8", errors="replace")
            logger.error(f"SmartParser: MinerU exited {proc.returncode}: {err[:500]}")
            return {"positions": [], "strategy": "mineru_failed", "error": err[:200]}
        md_files = list(output_dir.rglob("*.md"))
        if not md_files:
            return {"positions": [], "strategy": "mineru_no_output"}
        md_content = md_files[0].read_text(encoding="utf-8", errors="replace")
        positions = self._extract_positions_from_markdown(md_content)
        logger.info(f"SmartParser: MinerU extracted {len(positions)} positions")
        return {"positions": positions, "strategy": "mineru", "md_path": str(md_files[0])}

    def _extract_positions_from_markdown(self, md_content: str) -> list:
        """Uses module-level re (no redundant local import)."""
        positions = []
        rows = re.findall(r"<tr>(.*?)</tr>", md_content, re.DOTALL)
        for row in rows:
            cells = re.findall(r"<td[^>]*>(.*?)</td>", row, re.DOTALL)
            cells = [re.sub(r"<[^>]+>", "", c).strip() for c in cells]
            if len(cells) >= 2 and cells[0]:
                positions.append({
                    "code": cells[0] if len(cells) > 0 else "",
                    "name": cells[1] if len(cells) > 1 else "",
                    "unit": cells[2] if len(cells) > 2 else "",
                    "price": cells[-1] if len(cells) > 1 else "",
                })
        return positions

    def _detect_type(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext in (".xlsx", ".xls", ".xlsm"):
            return "excel"
        elif ext in (".xml",):
            return "xml"
        elif ext in (".docx",):
            return "docx"
        elif ext in (".csv", ".tsv"):
            return "csv"
        elif ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".gif", ".webp"):
            return "image"
        else:
            return "pdf"

    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        size_mb = _get_file_size_mb(path)
        return {
            "name": path.name,
            "size_mb": round(size_mb, 2),
            "type": self._detect_type(path),
            "strategy": "streaming" if size_mb > SIZE_THRESHOLD_MB else "standard",
            "mineru_enabled": ENABLE_MINERU,
        }
