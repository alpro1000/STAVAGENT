"""
Layer 1: Text Extraction — route to format-specific extractors.
"""

import logging
from pathlib import Path
from typing import Optional

from .models import ExtractedContent, FileFormat

logger = logging.getLogger(__name__)


async def extract_content(
    fmt: FileFormat,
    file_path: str,
    file_bytes: Optional[bytes] = None,
) -> ExtractedContent:
    """Extract text + tables from a file based on its format."""
    logger.info(f"[Extraction] Starting: format={fmt.value}, path={file_path}")

    if fmt == FileFormat.PDF:
        return await _extract_pdf(file_path)
    elif fmt in (FileFormat.XLSX, FileFormat.XLS):
        return _extract_excel(file_path)
    elif fmt == FileFormat.CSV:
        return _extract_csv(file_path)
    elif fmt == FileFormat.DOCX:
        return _extract_docx(file_path)
    elif fmt == FileFormat.XML:
        return _extract_xml(file_path)
    elif fmt in (FileFormat.JPEG, FileFormat.PNG, FileFormat.TIFF):
        return await _extract_image(file_path)
    elif fmt == FileFormat.DXF:
        return _extract_dxf(file_path)
    else:
        logger.warning(f"[Extraction] No extractor for {fmt.value}, returning empty")
        return ExtractedContent(parser_used="none")


# ── PDF ──────────────────────────────────────────────────────

async def _extract_pdf(file_path: str) -> ExtractedContent:
    """Extract text + tables from PDF using pdfplumber. Fallback to MinerU for scans."""
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber not installed")
        return ExtractedContent(parser_used="error:no_pdfplumber")

    text_parts: list[str] = []
    all_tables: list[list[list[str]]] = []
    page_count = 0

    try:
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)

                for table in page.extract_tables() or []:
                    cleaned = [
                        [str(cell) if cell is not None else "" for cell in row]
                        for row in table
                    ]
                    all_tables.append(cleaned)
    except Exception as e:
        logger.error(f"pdfplumber failed: {e}")
        return ExtractedContent(parser_used=f"error:pdfplumber:{e}")

    full_text = "\n".join(text_parts)
    chars = len(full_text)
    chars_per_page = chars / max(page_count, 1)

    # Quality check: if very low text density, likely a scan
    is_scan = chars_per_page < 200 and page_count > 0
    if is_scan:
        logger.info(
            f"[Extraction] Low text density ({chars_per_page:.0f} chars/page) — "
            f"likely scanned PDF, MinerU OCR recommended"
        )
        # Try MinerU async (if available)
        mineru_text = await _try_mineru_ocr(file_path)
        if mineru_text and len(mineru_text) > chars:
            full_text = mineru_text
            chars = len(full_text)

    quality = min(1.0, chars_per_page / 1000) if page_count > 0 else 0.0

    return ExtractedContent(
        text=full_text,
        tables=all_tables,
        metadata={"is_scan": is_scan, "chars_per_page": round(chars_per_page)},
        page_count=page_count,
        parser_used="pdfplumber" + ("+mineru" if is_scan and chars > 200 else ""),
        extraction_quality=round(quality, 2),
    )


async def _try_mineru_ocr(file_path: str) -> Optional[str]:
    """Try MinerU Cloud Run OCR. Returns text or None."""
    try:
        from app.parsers.mineru_client import MineruClient
        client = MineruClient()
        result = await client.parse_pdf(file_path)
        if result and result.get("text"):
            logger.info(f"[MinerU] OCR returned {len(result['text'])} chars")
            return result["text"]
    except Exception as e:
        logger.debug(f"[MinerU] OCR unavailable: {e}")
    return None


# ── Excel ────────────────────────────────────────────────────

def _extract_excel(file_path: str) -> ExtractedContent:
    """Extract text + tables from XLSX/XLS."""
    try:
        import openpyxl
    except ImportError:
        return ExtractedContent(parser_used="error:no_openpyxl")

    text_parts: list[str] = []
    all_tables: list[list[list[str]]] = []
    sheet_count = 0

    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheet_count = len(wb.sheetnames)

        for ws in wb.worksheets:
            sheet_rows: list[list[str]] = []
            for row in ws.iter_rows(values_only=True):
                str_row = [str(cell) if cell is not None else "" for cell in row]
                if any(c.strip() for c in str_row):
                    sheet_rows.append(str_row)
                    text_parts.append(" | ".join(c for c in str_row if c.strip()))

            if sheet_rows:
                all_tables.append(sheet_rows)

        wb.close()
    except Exception as e:
        logger.error(f"openpyxl failed: {e}")
        return ExtractedContent(parser_used=f"error:openpyxl:{e}")

    full_text = "\n".join(text_parts)
    return ExtractedContent(
        text=full_text,
        tables=all_tables,
        metadata={"sheet_count": sheet_count},
        page_count=sheet_count,
        parser_used="openpyxl",
        extraction_quality=0.9 if full_text else 0.0,
    )


# ── CSV ──────────────────────────────────────────────────────

def _extract_csv(file_path: str) -> ExtractedContent:
    """Extract from CSV/TSV."""
    import csv

    text_parts: list[str] = []
    all_rows: list[list[str]] = []

    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            sample = f.read(4096)
            f.seek(0)
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            reader = csv.reader(f, dialect)
            for row in reader:
                str_row = [str(c) for c in row]
                all_rows.append(str_row)
                text_parts.append(" | ".join(c for c in str_row if c.strip()))
    except Exception as e:
        logger.error(f"CSV parse failed: {e}")
        return ExtractedContent(parser_used=f"error:csv:{e}")

    return ExtractedContent(
        text="\n".join(text_parts),
        tables=[all_rows] if all_rows else [],
        page_count=1,
        parser_used="csv",
        extraction_quality=0.9,
    )


# ── DOCX ─────────────────────────────────────────────────────

def _extract_docx(file_path: str) -> ExtractedContent:
    """Extract text + tables from DOCX."""
    try:
        from docx import Document
    except ImportError:
        return ExtractedContent(parser_used="error:no_python_docx")

    text_parts: list[str] = []
    all_tables: list[list[list[str]]] = []

    try:
        doc = Document(file_path)

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            all_tables.append(rows)
    except Exception as e:
        logger.error(f"python-docx failed: {e}")
        return ExtractedContent(parser_used=f"error:docx:{e}")

    return ExtractedContent(
        text="\n".join(text_parts),
        tables=all_tables,
        metadata={"paragraphs": len(text_parts), "tables": len(all_tables)},
        page_count=max(1, len(text_parts) // 40),  # rough estimate
        parser_used="python-docx",
        extraction_quality=0.85,
    )


# ── XML ──────────────────────────────────────────────────────

def _extract_xml(file_path: str) -> ExtractedContent:
    """Extract text from XML."""
    try:
        from lxml import etree
    except ImportError:
        # Fallback to stdlib
        import xml.etree.ElementTree as etree  # type: ignore[no-redef]

    try:
        tree = etree.parse(file_path)
        root = tree.getroot()
        texts = [elem.text for elem in root.iter() if elem.text and elem.text.strip()]
        full_text = "\n".join(texts)
    except Exception as e:
        logger.error(f"XML parse failed: {e}")
        # Try as plain text
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                full_text = f.read()
        except Exception:
            return ExtractedContent(parser_used=f"error:xml:{e}")

    return ExtractedContent(
        text=full_text,
        page_count=1,
        parser_used="lxml",
        extraction_quality=0.8,
    )


# ── Image ────────────────────────────────────────────────────

async def _extract_image(file_path: str) -> ExtractedContent:
    """Extract text from image via OCR."""
    text = await _try_mineru_ocr(file_path)
    if text:
        return ExtractedContent(
            text=text,
            page_count=1,
            images_count=1,
            parser_used="mineru_ocr",
            extraction_quality=0.7,
        )
    return ExtractedContent(
        metadata={"note": "OCR not available"},
        page_count=1,
        images_count=1,
        parser_used="none",
        extraction_quality=0.0,
    )


# ── DXF ──────────────────────────────────────────────────────

def _extract_dxf(file_path: str) -> ExtractedContent:
    """Extract text layers from DXF."""
    try:
        import ezdxf
    except ImportError:
        return ExtractedContent(parser_used="error:no_ezdxf")

    text_parts: list[str] = []
    try:
        doc = ezdxf.readfile(file_path)
        msp = doc.modelspace()
        for entity in msp:
            if entity.dxftype() in ("TEXT", "MTEXT"):
                t = entity.dxf.text if hasattr(entity.dxf, "text") else ""
                if t and t.strip():
                    text_parts.append(t.strip())
    except Exception as e:
        logger.error(f"ezdxf failed: {e}")
        return ExtractedContent(parser_used=f"error:dxf:{e}")

    return ExtractedContent(
        text="\n".join(text_parts),
        metadata={"entity_texts": len(text_parts)},
        page_count=1,
        parser_used="ezdxf",
        extraction_quality=0.6,
    )
