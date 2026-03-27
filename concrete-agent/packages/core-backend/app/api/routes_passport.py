"""
Project Passport API Routes

Generate complete project passports from construction documents.

ENDPOINTS:
- POST /api/v1/passport/generate - Generate passport from file
- POST /api/v1/passport/generate-multiple - Generate from multiple files
- GET /api/v1/passport/{passport_id} - Get passport by ID
- GET /api/v1/passport/health - Health check

VERSION: 1.0.0 (2026-02-10)
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
import tempfile
import os

from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse

from app.services.document_processor import DocumentProcessor, process_document
from app.services.brief_summarizer import BriefDocumentSummarizer
from app.models.passport_schema import (
    ProjectPassport,
    PassportGenerationRequest,
    PassportGenerationResponse
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/passport", tags=["Project Passport"])

# In-memory storage for passports (TODO: Replace with database)
_passport_storage: Dict[str, ProjectPassport] = {}


# ============================================================================
# ENDPOINTS
# ============================================================================

@router.post("/generate")
async def generate_passport(
    file: UploadFile = File(..., description="Construction document (PDF, Excel, XML)"),
    project_name: str = Form(..., description="Project name"),
    enable_ai_enrichment: bool = Form(True, description="Enable AI enrichment (Layer 3)"),
    preferred_model: Optional[str] = Form(None, description="Preferred AI model: gemini, claude-sonnet, claude-haiku, openai, openai-mini, perplexity, vertex-ai-gemini, vertex-ai-search, auto"),
    project_id: Optional[str] = Form(None, description="Optional project ID"),
    vertex_service_account: Optional[str] = Form(None, description="Vertex AI service account ID (e.g. stavagent-vertex-search@...)"),
    llm_provider: Optional[str] = Form(None, description="LLM provider hint: vertex-ai, google, anthropic, openai"),
    analysis_mode: Optional[str] = Form("adaptive_extraction", description="Analysis mode: adaptive_extraction (full passport) or summary_only (adaptive topics)")
):
    """
    Generate project passport or adaptive summary from a document.

    **Modes:**
    - `adaptive_extraction` (default): Full 3-layer passport (structured JSON, 50+ fields)
    - `summary_only`: Adaptive topic-based summary (dynamic topics, any document type)

    **3-Layer Architecture (adaptive_extraction):**
    - Layer 1: Document parsing (MinerU/SmartParser) - 1-3s
    - Layer 2: Regex extraction (deterministic facts) - <100ms
    - Layer 3: AI enrichment (Claude/Gemini, optional) - 3-5s

    **Adaptive Summary (summary_only):**
    NotebookLM-inspired 2-step approach:
    - Step 1 (INDEX): Discover all topics in the document
    - Step 2 (EXPLAIN): Deep-dive each topic with facts and numbers
    - Works with ANY document type (not limited to construction)

    **Example:**
    ```bash
    # Full passport
    curl -X POST ".../api/v1/passport/generate" \\
      -F "file=@technicka_zprava.pdf" \\
      -F "project_name=Polyfunkční dům Praha 5"

    # Adaptive summary
    curl -X POST ".../api/v1/passport/generate" \\
      -F "file=@document.pdf" \\
      -F "project_name=Any Document" \\
      -F "analysis_mode=summary_only"
    ```
    """
    logger.info(f"Generating passport: {project_name}, AI: {enable_ai_enrichment}, mode: {analysis_mode}")

    # Validate file type
    allowed_extensions = [
        '.pdf', '.xlsx', '.xls', '.xml', '.docx', '.csv',
        '.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif', '.webp',
        '.dxf', '.dwg',
    ]
    file_ext = Path(file.filename).suffix.lower()

    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file_ext}. Allowed: {allowed_extensions}"
        )

    # Save uploaded file to temp location
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, file.filename)

    try:
        # Write uploaded file
        with open(temp_file_path, "wb") as f:
            content = await file.read()
            f.write(content)

        logger.info(f"Saved file to: {temp_file_path} ({len(content)} bytes)")

        # ── v5.0: Try universal parser for XLSX/structured formats ──
        # Returns ParsedDocument with positions, chapters, SO — alongside Passport
        soupis_data = None
        if file_ext in ('.xlsx', '.xlsm', '.xls'):
            try:
                from app.parsers.universal_parser import parse_any
                parsed_doc = parse_any(temp_file_path)
                if parsed_doc.positions_count > 0:
                    soupis_data = {
                        "format": parsed_doc.source_format.value,
                        "project_name": parsed_doc.project_name,
                        "project_id": parsed_doc.project_id,
                        "positions_count": parsed_doc.positions_count,
                        "coverage_pct": parsed_doc.coverage_pct,
                        "so_count": len(parsed_doc.stavebni_objekty),
                        "stavebni_objekty": [
                            {
                                "so_id": so.so_id,
                                "so_name": so.so_name,
                                "chapters_count": len(so.chapters),
                                "positions_count": sum(len(ch.positions) for ch in so.chapters),
                                "chapters": [
                                    {
                                        "code": ch.code,
                                        "name": ch.name,
                                        "positions": [
                                            {
                                                "pc": p.pc,
                                                "code": p.code,
                                                "description": p.description,
                                                "unit": p.unit,
                                                "quantity": float(p.quantity) if p.quantity else None,
                                                "unit_price": float(p.unit_price) if p.unit_price else None,
                                                "total_price": float(p.total_price) if p.total_price else None,
                                                "price_source": p.price_source,
                                                "specification": p.specification[:200] if p.specification else None,
                                                "vv_lines_count": len(p.vv_lines),
                                                "url": p.url,
                                            }
                                            for p in ch.positions
                                        ],
                                    }
                                    for ch in so.chapters
                                ],
                            }
                            for so in parsed_doc.stavebni_objekty
                        ],
                        "warnings": parsed_doc.parser_warnings[:10],
                    }
                    logger.info(
                        f"Universal parser: {parsed_doc.positions_count} positions, "
                        f"format={parsed_doc.source_format.value}"
                    )
            except Exception as e:
                logger.warning(f"Universal parser failed (fallback to SmartParser): {e}")

        # Determine effective model
        effective_model = preferred_model
        if not effective_model and llm_provider:
            provider_defaults = {
                "vertex-ai": "vertex-ai-gemini",
                "google": "gemini",
                "anthropic": "claude-sonnet",
                "openai": "openai",
            }
            effective_model = provider_defaults.get(llm_provider)

        # === SUMMARY_ONLY MODE: Adaptive topic-based summary ===
        if analysis_mode == "summary_only":
            return await _generate_adaptive_summary(
                temp_file_path, project_name, effective_model or "gemini"
            )

        # === ADAPTIVE_EXTRACTION MODE: Full passport ===
        processor = DocumentProcessor(
            preferred_model=effective_model,
            vertex_service_account=vertex_service_account,
        )
        response = await processor.process(
            file_path=temp_file_path,
            project_name=project_name,
            enable_ai_enrichment=enable_ai_enrichment,
            preferred_model=effective_model,
            project_id=project_id
        )

        if response.success and response.passport:
            # Store passport in memory
            passport_id = response.passport.passport_id
            _passport_storage[passport_id] = response.passport

            logger.info(f"Passport generated: {passport_id}, time: {response.processing_time_ms}ms")

        # v5.0: Attach soupis data if available (from universal parser)
        if soupis_data:
            resp_dict = response.model_dump() if hasattr(response, 'model_dump') else response.dict()
            resp_dict["soupis_praci"] = soupis_data
            return JSONResponse(content=resp_dict)

        return response

    except Exception as e:
        logger.error(f"Failed to generate passport: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

    finally:
        # Cleanup temp file
        try:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            if os.path.exists(temp_dir):
                os.rmdir(temp_dir)
        except Exception as e:
            logger.warning(f"Failed to cleanup temp files: {e}")


async def _generate_adaptive_summary(
    file_path: str, project_name: str, preferred_model: str
) -> JSONResponse:
    """
    Generate adaptive topic-based summary (summary_only mode).

    Extracts text from document, then uses NotebookLM-inspired
    2-step approach (INDEX → EXPLAIN) for universal document analysis.
    """
    import time as _time
    start = _time.time()

    # Extract text from document — read ALL pages for large documents
    document_text = ""
    total_pages = 0
    file_ext = Path(file_path).suffix.lower()

    if file_ext == '.pdf':
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"PDF has {total_pages} pages — reading all")
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        document_text += page_text + "\n"
                logger.info(f"Extracted {len(document_text)} chars from {total_pages} pages")
        except Exception as e:
            logger.warning(f"Failed to extract PDF text: {e}")

    elif file_ext == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            document_text = "\n".join(paragraphs)
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    document_text += "\n" + " | ".join(cells)
            total_pages = max(1, len(paragraphs) // 30)
            logger.info(f"DOCX extracted: {len(document_text)} chars, ~{total_pages} pages")
        except Exception as e:
            logger.warning(f"Failed to extract DOCX text: {e}")

    elif file_ext == '.csv':
        try:
            import pandas as pd
            df = pd.read_csv(file_path, encoding="utf-8", on_bad_lines="skip")
            document_text = df.to_string()
            total_pages = 1
            logger.info(f"CSV extracted: {len(document_text)} chars, {len(df)} rows")
        except Exception as e:
            logger.warning(f"Failed to extract CSV: {e}")

    if not document_text:
        # Fallback: try SmartParser for Excel/XML
        try:
            from app.parsers.smart_parser import SmartParser
            parser = SmartParser()
            parsed = parser.parse(Path(file_path))
            if 'positions' in parsed:
                for pos in parsed['positions']:
                    if 'popis' in pos:
                        document_text += pos['popis'] + " "
            if 'text' in parsed:
                document_text = parsed['text']
        except Exception as e:
            logger.warning(f"SmartParser fallback failed: {e}")

    if not document_text:
        raise HTTPException(status_code=400, detail="Failed to extract text from document")

    # Classify document (3-tier with AI fallback)
    from app.services.document_classifier import classify_document_async
    from app.services.passport_enricher import PassportEnricher
    _enricher = PassportEnricher(preferred_model=preferred_model)
    classification = await classify_document_async(
        Path(file_path).name, document_text, llm_call=_enricher._call_llm
    )

    # Generate adaptive summary
    summarizer = BriefDocumentSummarizer(preferred_model=preferred_model)
    result = await summarizer.summarize(
        document_text=document_text,
        language="cs",
        max_chars=0  # use model's native context limit (was 8000 — too small for 40-50 page docs)
    )

    # Wrap in passport-compatible response format
    # so frontend can detect and render appropriately
    return JSONResponse(content={
        "success": True,
        "analysis_mode": "summary_only",
        "format": "adaptive_v2",
        "project_name": project_name,
        "adaptive_summary": result,
        "classification": classification.model_dump(),
        # Minimal passport stub for backward compatibility
        "passport": {
            "passport_id": f"summary_{int(_time.time())}",
            "project_name": project_name,
            "generated_at": datetime.now().isoformat(),
            "source_documents": [Path(file_path).name],
            "description": result.get("summary", ""),
            "document_type": result.get("document_type", ""),
            "topics": result.get("topics", []),
            "warnings": result.get("warnings", []),
            # Empty arrays for structured fields (not applicable in summary mode)
            "concrete_specifications": [],
            "reinforcement": [],
            "quantities": [],
            "dimensions": None,
            "special_requirements": [],
            "objects": [],
            "structure_type": None,
            "location": None,
            "timeline": None,
            "stakeholders": [],
            "risks": [],
            "technical_highlights": [],
            "extraction_stats": {},
            "processing_time_ms": result.get("processing_time_ms", 0),
            "layer_breakdown": {},
        },
        "processing_time_ms": result.get("processing_time_ms", 0),
        "metadata": {
            "file_name": Path(file_path).name,
            "processing_time_seconds": round((_time.time() - start), 2),
            "parser_used": "AdaptiveSummarizer",
            "extraction_method": "NotebookLM-inspired INDEX→EXPLAIN",
            "ai_model_used": result.get("model_used", preferred_model),
            "total_confidence": 0.85,
        },
        "statistics": {
            "total_concrete_m3": 0,
            "total_reinforcement_t": 0,
            "unique_concrete_classes": 0,
            "unique_steel_grades": 0,
            "deterministic_fields": 0,
            "ai_enriched_fields": len(result.get("topics", [])),
        },
    })


@router.post("/summarize")
async def summarize_document_brief(
    file: UploadFile = File(..., description="Construction document (PDF, Excel, XML)"),
    language: str = Form("cs", description="Summary language (cs, en)"),
    preferred_model: Optional[str] = Form("gemini", description="AI model: gemini (FREE), claude-haiku, openai-mini")
):
    """
    Generate BRIEF TEXT SUMMARY of document (5-10 sentences).
    
    **Difference from /generate:**
    - /generate: Full structured passport (JSON, 50+ fields, 4-8s)
    - /summarize: Brief text summary (plain text, 5-10 sentences, 2-3s)
    
    **Use cases:**
    - Quick document overview
    - Email notifications
    - Dashboard previews
    
    **Performance:**
    - Processing time: 2-3 seconds (vs 4-8s for full passport)
    - Prompt length: 2K chars (vs 5K for passport)
    - Output: Plain text (vs complex JSON)
    
    **Example:**
    ```bash
    curl -X POST "http://localhost:8000/api/v1/passport/summarize" \\
      -F "file=@technicka_zprava.pdf" \\
      -F "language=cs" \\
      -F "preferred_model=gemini"
    ```
    
    **Response:**
    ```json
    {
      "summary": "Projekt: Most přes Chrudimku km 15.2\nTyp: Silniční most...\n",
      "processing_time_ms": 2500,
      "chars_processed": 2000,
      "model_used": "gemini"
    }
    ```
    """
    logger.info(f"Generating brief summary: {file.filename}, language: {language}")
    
    # Save uploaded file to temp location
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, file.filename)
    
    try:
        # Write uploaded file
        with open(temp_file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        logger.info(f"Saved file to: {temp_file_path} ({len(content)} bytes)")
        
        # Extract text from document
        from app.parsers.smart_parser import SmartParser
        parser = SmartParser()
        parsed = parser.parse(Path(temp_file_path))
        
        # Get text content
        document_text = ""
        
        # For PDFs: extract full text with pdfplumber
        if Path(temp_file_path).suffix.lower() == '.pdf':
            try:
                import pdfplumber
                with pdfplumber.open(temp_file_path) as pdf:
                    for page in pdf.pages[:5]:  # First 5 pages only
                        page_text = page.extract_text()
                        if page_text:
                            document_text += page_text + "\n"
            except Exception as e:
                logger.warning(f"Failed to extract PDF text: {e}")
        
        # Fallback to parsed data
        if not document_text and 'positions' in parsed:
            for pos in parsed['positions']:
                if 'popis' in pos:
                    document_text += pos['popis'] + " "
        
        if not document_text:
            raise HTTPException(
                status_code=400,
                detail="Failed to extract text from document"
            )
        
        # Generate brief summary
        summarizer = BriefDocumentSummarizer(preferred_model=preferred_model)
        result = await summarizer.summarize(
            document_text=document_text,
            language=language,
            max_chars=2000  # Only first 2K chars
        )
        
        logger.info(f"Brief summary generated: {result['processing_time_ms']}ms")
        
        return result
    
    except Exception as e:
        logger.error(f"Failed to generate brief summary: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        # Cleanup temp file
        try:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            if os.path.exists(temp_dir):
                os.rmdir(temp_dir)
        except Exception as e:
            logger.warning(f"Failed to cleanup temp files: {e}")


@router.post("/generate-from-path")
async def generate_passport_from_path(
    request: PassportGenerationRequest
) -> PassportGenerationResponse:
    """
    Generate passport from existing file path.

    Useful when file is already uploaded to server.

    **Request Body:**
    ```json
    {
      "project_id": "proj_12345",
      "file_paths": ["/data/projects/proj_12345/technicka_zprava.pdf"],
      "enable_ai_enrichment": true,
      "language": "cs"
    }
    ```
    """
    logger.info(f"Generating passport from path: {request.file_paths}")

    if not request.file_paths:
        raise HTTPException(status_code=400, detail="No file paths provided")

    # Process first file (or all files and merge)
    file_path = request.file_paths[0]

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail=f"File not found: {file_path}")

    try:
        response = await process_document(
            file_path=file_path,
            project_name=request.project_id,  # Use project_id as name for now
            enable_ai=request.enable_ai_enrichment,
            project_id=request.project_id
        )

        if response.success and response.passport:
            # Store passport
            passport_id = response.passport.passport_id
            _passport_storage[passport_id] = response.passport

        return response

    except Exception as e:
        logger.error(f"Failed to generate passport: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{passport_id}", response_model=ProjectPassport)
async def get_passport(passport_id: str):
    """
    Get passport by ID.

    **Example:**
    ```bash
    curl "http://localhost:8000/api/v1/passport/passport_abc123"
    ```
    """
    if passport_id not in _passport_storage:
        raise HTTPException(
            status_code=404,
            detail=f"Passport not found: {passport_id}"
        )

    return _passport_storage[passport_id]


@router.get("/list/all")
async def list_passports():
    """
    List all generated passports (IDs only).

    **Example:**
    ```bash
    curl "http://localhost:8000/api/v1/passport/list/all"
    ```
    """
    return {
        "total": len(_passport_storage),
        "passports": [
            {
                "passport_id": p.passport_id,
                "project_name": p.project_name,
                "generated_at": p.generated_at.isoformat(),
                "source_documents": p.source_documents
            }
            for p in _passport_storage.values()
        ]
    }


@router.delete("/{passport_id}")
async def delete_passport(passport_id: str):
    """
    Delete passport by ID.

    **Example:**
    ```bash
    curl -X DELETE "http://localhost:8000/api/v1/passport/passport_abc123"
    ```
    """
    if passport_id not in _passport_storage:
        raise HTTPException(
            status_code=404,
            detail=f"Passport not found: {passport_id}"
        )

    del _passport_storage[passport_id]

    return {"message": f"Passport {passport_id} deleted"}


@router.get("/health")
async def health_check():
    """
    Health check for passport service.

    Returns:
    - Service status
    - Number of passports in memory
    - Layer availability (Parser, Regex, AI)
    """
    # Check if services are available
    try:
        processor = DocumentProcessor()
        has_ai = processor.enricher.gemini_model is not None or processor.enricher.claude_client is not None
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        has_ai = False

    return {
        "status": "healthy",
        "passports_in_memory": len(_passport_storage),
        "layers": {
            "layer1_parser": "available",
            "layer2_regex": "available",
            "layer3_ai": "available" if has_ai else "unavailable (no API keys)"
        },
        "timestamp": datetime.now().isoformat()
    }


# ============================================================================
# EXPORT ENDPOINTS
# ============================================================================

@router.get("/{passport_id}/export/json")
async def export_passport_json(passport_id: str):
    """
    Export passport as JSON.
    """
    if passport_id not in _passport_storage:
        raise HTTPException(
            status_code=404,
            detail=f"Passport not found: {passport_id}"
        )

    passport = _passport_storage[passport_id]

    return JSONResponse(
        content=passport.model_dump(mode='json'),
        headers={
            "Content-Disposition": f"attachment; filename={passport_id}.json"
        }
    )


@router.get("/{passport_id}/summary")
async def get_passport_summary(passport_id: str):
    """
    Get compact summary of passport (key facts only).
    """
    if passport_id not in _passport_storage:
        raise HTTPException(
            status_code=404,
            detail=f"Passport not found: {passport_id}"
        )

    passport = _passport_storage[passport_id]

    # Build summary
    summary = {
        "passport_id": passport.passport_id,
        "project_name": passport.project_name,
        "generated_at": passport.generated_at.isoformat(),

        # Key facts
        "concrete_classes": [
            {
                "class": spec.concrete_class,
                "exposure": [e.value for e in spec.exposure_classes]
            }
            for spec in passport.concrete_specifications
        ],

        "steel_grades": [
            steel.steel_grade.value
            for steel in passport.reinforcement
        ],

        "dimensions": {
            "floors_underground": passport.dimensions.floors_underground if passport.dimensions else None,
            "floors_above_ground": passport.dimensions.floors_above_ground if passport.dimensions else None,
            "height_m": passport.dimensions.height_m if passport.dimensions else None
        },

        "special_requirements": [
            req.requirement_type
            for req in passport.special_requirements
        ],

        "total_quantities": {
            "volume_m3": sum(q.volume_m3 or 0 for q in passport.quantities),
            "area_m2": sum(q.area_m2 or 0 for q in passport.quantities),
            "mass_tons": sum(q.mass_tons or 0 for q in passport.quantities)
        },

        # AI enrichments
        "structure_type": passport.structure_type.value if passport.structure_type else None,
        "risks_count": len(passport.risks),
        "stakeholders_count": len(passport.stakeholders),

        # Processing stats
        "processing_time_ms": passport.processing_time_ms,
        "extraction_stats": passport.extraction_stats
    }

    return summary




# ============================================================================
# v3: MULTI-DOCUMENT PROJECT PROCESSING (SO-based merge)
# ============================================================================

@router.post("/process-project")
async def process_project(
    files: list[UploadFile] = File(..., description="Multiple construction documents"),
    project_name: str = Form(..., description="Project name"),
    enable_ai_enrichment: bool = Form(True, description="Enable AI enrichment (Layer 3)"),
    preferred_model: Optional[str] = Form(None, description="Preferred AI model"),
    project_id: Optional[str] = Form(None, description="Optional project ID"),
):
    """
    v3: Process multiple documents with SO-based grouping and merge.

    Groups files by SO code from filenames, processes each through the
    3-layer pipeline, and merges results per-SO with contradiction detection.

    **Filename conventions:**
    - `201_01_TZ.pdf` → SO 201 (dílčí TZ)
    - `SO_202_situace.pdf` → SO 202 (výkres)
    - `technicka_zprava.pdf` → SO 000 (project-level)

    **Returns:**
    - `merged_sos`: List of MergedSO objects with bridge_params, gtp, tender
    - `contradictions`: List of detected contradictions across documents
    - `file_groups`: SOFileGroup list with coverage per SO
    - Standard passport fields for backward compatibility

    **Example:**
    ```bash
    curl -X POST ".../api/v1/passport/process-project" \\
      -F "files=@201_01_TZ.pdf" \\
      -F "files=@201_02_situace.pdf" \\
      -F "files=@202_01_TZ.pdf" \\
      -F "files=@zadavaci_dokumentace.pdf" \\
      -F "project_name=Most přes Chrudimku"
    ```
    """
    import time as _time
    from app.services.file_grouper import group_files_by_so, get_coverage_report
    from app.services.so_merger import merge_so_group

    start = _time.time()
    logger.info(f"Processing project: {project_name}, {len(files)} files")

    # Validate files
    allowed_extensions = [
        '.pdf', '.xlsx', '.xls', '.xml', '.docx', '.csv',
        '.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif', '.webp',
        '.dxf', '.dwg',
    ]
    temp_dir = tempfile.mkdtemp()
    saved_paths: list[str] = []

    try:
        # Save all uploaded files
        for f in files:
            ext = Path(f.filename).suffix.lower()
            if ext not in allowed_extensions:
                logger.warning(f"Skipping unsupported file: {f.filename}")
                continue
            temp_path = os.path.join(temp_dir, f.filename)
            with open(temp_path, "wb") as out:
                content = await f.read()
                out.write(content)
            saved_paths.append(temp_path)
            logger.info(f"Saved: {f.filename} ({len(content)} bytes)")

        if not saved_paths:
            raise HTTPException(status_code=400, detail="No valid files provided")

        # v5.0: Try universal parser for XLSX files (soupis prací extraction)
        all_soupis = []
        for sp in saved_paths:
            if Path(sp).suffix.lower() in ('.xlsx', '.xlsm', '.xls'):
                try:
                    from app.parsers.universal_parser import parse_any
                    pd = parse_any(sp)
                    if pd.positions_count > 0:
                        all_soupis.append({
                            "file": Path(sp).name,
                            "format": pd.source_format.value,
                            "positions_count": pd.positions_count,
                            "coverage_pct": pd.coverage_pct,
                            "so_count": len(pd.stavebni_objekty),
                            "stavebni_objekty": [
                                {
                                    "so_id": so.so_id,
                                    "so_name": so.so_name,
                                    "chapters": [
                                        {
                                            "code": ch.code,
                                            "name": ch.name,
                                            "positions_count": len(ch.positions),
                                            "positions": [
                                                {
                                                    "pc": p.pc,
                                                    "code": p.code,
                                                    "description": p.description,
                                                    "unit": p.unit,
                                                    "quantity": float(p.quantity) if p.quantity else None,
                                                    "unit_price": float(p.unit_price) if p.unit_price else None,
                                                    "total_price": float(p.total_price) if p.total_price else None,
                                                    "price_source": p.price_source,
                                                }
                                                for p in ch.positions
                                            ],
                                        }
                                        for ch in so.chapters
                                    ],
                                }
                                for so in pd.stavebni_objekty
                            ],
                        })
                        logger.info(f"Universal parser: {Path(sp).name} → {pd.positions_count} positions")
                except Exception as e:
                    logger.warning(f"Universal parser skipped for {Path(sp).name}: {e}")

        # Step 1: Group files by SO code
        file_groups = group_files_by_so(saved_paths)
        coverage_report = get_coverage_report(file_groups)

        logger.info(
            f"File grouping: {len(saved_paths)} files → "
            f"{len(file_groups)} SO groups"
        )

        # Step 2: Process each file through the pipeline
        processor = DocumentProcessor(preferred_model=preferred_model)
        file_results: Dict[str, Dict[str, Any]] = {}  # filename → results

        for path in saved_paths:
            filename = Path(path).name
            logger.info(f"Processing: {filename}")

            try:
                response = await processor.process(
                    file_path=path,
                    project_name=project_name,
                    enable_ai_enrichment=enable_ai_enrichment,
                    preferred_model=preferred_model,
                    project_id=project_id,
                )

                result_dict: Dict[str, Any] = {
                    "filename": filename,
                    "success": response.success,
                    "classification": response.classification,
                }

                if response.success:
                    result_dict["technical"] = (
                        response.technical.model_dump() if response.technical else None
                    )
                    result_dict["bill_of_quantities"] = (
                        response.bill_of_quantities.model_dump()
                        if response.bill_of_quantities else None
                    )
                    result_dict["tender_conditions"] = (
                        response.tender_conditions.model_dump()
                        if response.tender_conditions else None
                    )
                    result_dict["schedule"] = (
                        response.schedule.model_dump() if response.schedule else None
                    )
                    # v3 fields
                    result_dict["tender"] = (
                        response.tender.model_dump() if response.tender else None
                    )
                    result_dict["gtp"] = (
                        response.gtp.model_dump() if response.gtp else None
                    )
                    result_dict["bridge_params"] = (
                        response.bridge_params.model_dump()
                        if response.bridge_params else None
                    )
                    result_dict["passport"] = (
                        response.passport.model_dump() if response.passport else None
                    )
                    # v3.1.1: Enhanced classification metadata
                    enhanced_meta = getattr(processor, '_last_enhanced_metadata', None)
                    if enhanced_meta:
                        result_dict["section_ids"] = enhanced_meta.get("section_ids", [])
                        result_dict["construction_type"] = enhanced_meta.get("construction_type")
                        result_dict["is_non_construction"] = enhanced_meta.get("is_non_construction", False)

                file_results[filename] = result_dict

            except Exception as e:
                logger.error(f"Failed to process {filename}: {e}")
                file_results[filename] = {
                    "filename": filename,
                    "success": False,
                    "error": str(e),
                }

        # Step 3: Merge results per SO group
        merged_sos = []
        all_contradictions = []

        for group in file_groups:
            group_results = []
            for so_file in group.files:
                fr = file_results.get(so_file.filename)
                if fr and fr.get("success"):
                    group_results.append(fr)

            if group_results:
                merged = merge_so_group(group.so_code, group_results)
                merged_sos.append(merged)
                all_contradictions.extend(merged.contradictions)

        total_time = int((_time.time() - start) * 1000)

        logger.info(
            f"Project processing complete: {total_time}ms, "
            f"{len(merged_sos)} merged SOs, "
            f"{len(all_contradictions)} contradictions"
        )

        # Build response
        return JSONResponse(content={
            "success": True,
            "project_name": project_name,
            "processing_time_ms": total_time,
            "file_count": len(saved_paths),

            # v3: Multi-document results
            "merged_sos": [so.model_dump(mode="json") for so in merged_sos],
            "contradictions": [c.model_dump(mode="json") for c in all_contradictions],
            "file_groups": [g.model_dump(mode="json") for g in file_groups],
            "coverage_report": coverage_report,

            # Per-file results (for debugging / detail view)
            "file_results": {
                fn: {
                    "success": fr.get("success"),
                    "classification": (
                        fr["classification"].model_dump(mode="json")
                        if fr.get("classification") and hasattr(fr["classification"], "model_dump")
                        else fr.get("classification")
                    ),
                }
                for fn, fr in file_results.items()
            },

            # Summary stats
            "statistics": {
                "total_files": len(saved_paths),
                "successful_files": sum(
                    1 for fr in file_results.values() if fr.get("success")
                ),
                "failed_files": sum(
                    1 for fr in file_results.values() if not fr.get("success")
                ),
                "so_groups": len(merged_sos),
                "contradictions_total": len(all_contradictions),
                "contradictions_critical": sum(
                    1 for c in all_contradictions if c.severity == "critical"
                ),
            },

            # v5.0: Soupis prací (from universal parser, XLSX files only)
            "soupis_praci": all_soupis if all_soupis else None,
        })

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Project processing failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

    finally:
        # Cleanup temp files
        import shutil
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception as e:
            logger.warning(f"Failed to cleanup temp dir: {e}")
